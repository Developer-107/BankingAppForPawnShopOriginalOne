import os
from utils import get_conn, office_mob_number
from datetime import datetime
import win32com.client
from docx import Document
from PyQt5.QtCore import QSize, Qt, QDateTime
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import QWidget, QGridLayout, QLabel, QLineEdit, QToolButton, QHBoxLayout, QMessageBox, QComboBox, \
    QCompleter

from utils import resource_path


class AddWindow(QWidget):
    def __init__(self, organisation, name_of_user):
        super().__init__()
        self.setWindowTitle("დამატება")
        self.setWindowIcon(QIcon(resource_path("Icons/add_data.png")))
        self.resize(1400, 500)

        self.name_of_user = name_of_user
        self.organisation = organisation

        layout = QGridLayout()

        take_date = QLabel("მიღების თარიღი")
        name_surname = QLabel("სახელი და გვარი:")
        id_number = QLabel("პირადი ნომერი:")
        tel_number = QLabel("საკონტაქტო ტელეფონი:")
        item_name = QLabel("ნივთის დასახელება:")
        model = QLabel("მოდელი:")
        imei_sn = QLabel("IMEI:")
        choose_the_type = QLabel("აირჩიეთ ტიპი:")

        self.take_date_box = QLabel(f"{QDateTime.currentDateTime().toString('yyyy-MM-dd HH:mm:ss')}")
        self.take_date_box.setStyleSheet("border: 1px solid gray; padding: 5px;")

        self.name_surname_box = QLineEdit()
        self.name_surname_box.setPlaceholderText("სახელი და გვარი")

        self.id_number_box = QLineEdit()
        self.id_number_box.setPlaceholderText("პირადი ნომერი")

        self.tel_number_box = QLineEdit()
        self.tel_number_box.setPlaceholderText("საკონტაქტო ტელეფონი")

        self.item_name_box = QLineEdit()
        self.item_name_box.setPlaceholderText("ნივთის დასახელება")

        self.model_box = QLineEdit()
        self.model_box.setPlaceholderText("მოდელი")

        self.imei_sn_box = QLineEdit()
        self.imei_sn_box.setPlaceholderText("IMEI")

        self.choose_the_type_box = QLabel("დატოვებით")
        self.choose_the_type_box.setStyleSheet("border: 1px solid gray; padding: 5px;")

        layout.addWidget(take_date, 0, 0)
        layout.addWidget(name_surname, 1, 0)
        layout.addWidget(id_number, 2, 0)
        layout.addWidget(tel_number, 3, 0)
        layout.addWidget(item_name, 4, 0)
        layout.addWidget(model, 5, 0)
        layout.addWidget(imei_sn, 6, 0)
        layout.addWidget(choose_the_type, 7, 0)

        layout.addWidget(self.take_date_box, 0, 1)
        layout.addWidget(self.name_surname_box, 1, 1)
        layout.addWidget(self.id_number_box, 2, 1)
        layout.addWidget(self.tel_number_box, 3, 1)
        layout.addWidget(self.item_name_box, 4, 1)
        layout.addWidget(self.model_box, 5, 1)
        layout.addWidget(self.imei_sn_box, 6, 1)
        layout.addWidget(self.choose_the_type_box, 7, 1)


        # Set up autocomplete and autofill
        self.name_list = self.load_name_list_from_db()
        self.completer = QCompleter(self.name_list)
        self.completer.setCaseSensitivity(Qt.CaseInsensitive)
        self.name_surname_box.setCompleter(self.completer)
        self.name_surname_box.editingFinished.connect(self.autofill_id_from_name)


        trusted_person = QLabel("მინდობილობა:")
        comment = QLabel("კომენტარი:")
        given_money = QLabel("გაცემული თანხა:")
        percent = QLabel("პროცენტის მითითება:")
        day_quantity = QLabel("დღის მითითება:")

        self.trusted_person_box = QLineEdit()
        self.trusted_person_box.setPlaceholderText("მინდობილობის გაკეთება ხდება კლიენტის სურვილით")
        self.trusted_person_box.setFixedHeight(70)

        self.comment_box = QLineEdit()
        self.comment_box.setPlaceholderText("კომენტარის გარეშე")
        self.comment_box.setFixedHeight(100)

        self.given_money_box = QLineEdit()
        self.given_money_box.setPlaceholderText("თანხა რამდენიც უნდა გასცეთ")


        # --- Percent ComboBox ---
        self.percent_box = QComboBox()
        self.percent_box.addItems(["2.5", "5", "10", "15"])
        self.percent_box.setCurrentText("10")
        self.percent_box.setStyleSheet("padding: 5px; font-size: 14px;")

        # --- Day Quantity ComboBox ---
        self.day_quantity_box = QComboBox()
        self.day_quantity_box.addItems(["10", "15", "30"])
        self.day_quantity_box.setCurrentText("10")
        self.day_quantity_box.setStyleSheet("padding: 5px; font-size: 14px;")

        layout.addWidget(trusted_person, 0, 2)
        layout.addWidget(comment, 1, 2)
        layout.addWidget(given_money, 2, 2)
        layout.addWidget(percent, 3, 2)
        layout.addWidget(day_quantity, 4, 2)

        layout.addWidget(self.trusted_person_box, 0, 3)
        layout.addWidget(self.comment_box, 1, 3)
        layout.addWidget(self.given_money_box, 2, 3)
        layout.addWidget(self.percent_box, 3, 3)
        layout.addWidget(self.day_quantity_box, 4, 3)


        save_button = QToolButton()
        save_button.setText(" შენახვა ")
        save_button.setIcon(QIcon(resource_path("Icons/save_icon.png")))
        save_button.setIconSize(QSize(35, 35))
        save_button.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        save_button.setStyleSheet("font-size: 16px;")
        save_button.clicked.connect(self.save_to_sql)

        close_window = QToolButton()
        close_window.setText(" დახურვა ")
        close_window.setIcon(QIcon(resource_path("Icons/cancel_icon.png")))
        close_window.setIconSize(QSize(35, 35))
        close_window.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        close_window.setStyleSheet("font-size: 16px;")
        close_window.clicked.connect(self.close)

        # Create horizontal layout for both buttons
        button_container = QWidget()
        button_layout = QHBoxLayout()
        button_layout.setContentsMargins(0, 0, 0, 0)  # remove padding
        button_layout.setSpacing(10)  # space between buttons

        button_layout.addWidget(save_button)
        button_layout.addWidget(close_window)
        button_container.setLayout(button_layout)

        # Add to main layout at row 5, column 3
        layout.addWidget(button_container, 7, 3)






        # --------------------------------------------Layout-----------------------------------------------------
        self.setLayout(layout)

    def load_name_list_from_db(self):
        try:
            conn = get_conn()
            cursor = conn.cursor()
            cursor.execute("SELECT DISTINCT name_surname FROM active_contracts")
            names = [row[0] for row in cursor.fetchall()]
            conn.close()
            return names
        except Exception as e:
            print("Error loading names:", e)
            return []



    def autofill_id_from_name(self):
        name = self.name_surname_box.text()
        if not name:
            return
        try:
            conn = get_conn()
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id_number FROM active_contracts 
                WHERE name_surname = %s 
                ORDER BY date DESC LIMIT 1
            """, (name,))
            result = cursor.fetchone()
            conn.close()
            if result:
                self.id_number_box.setText(str(result[0]))
            else:
                self.id_number_box.clear()
        except Exception as e:
            print("Autofill error:", e)


    def save_to_sql(self):
        try:
            conn = get_conn()
            cursor = conn.cursor()

            cursor.execute("""
                INSERT INTO active_contracts (
                    date, name_surname, id_number, tel_number,
                    item_name, model, imei, type,
                    trusted_person, comment, given_money,
                    percent, day_quantity, added_percents
                ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (
                self.take_date_box.text(),
                self.name_surname_box.text(),
                self.id_number_box.text(),
                self.tel_number_box.text(),
                self.item_name_box.text(),
                self.model_box.text(),
                self.imei_sn_box.text(),
                self.choose_the_type_box.text(),
                self.trusted_person_box.text(),
                self.comment_box.text(),
                self.given_money_box.text(),
                float(self.percent_box.currentText()),
                int(self.day_quantity_box.currentText()),
                float(int(self.given_money_box.text()) * float(self.percent_box.currentText()) / 100)
            ))

            contract_id = cursor.fetchone()[0]


            status_for_given_principle = "გაცემული ძირი თანხა"

            cursor.execute("""
                            INSERT INTO given_and_additional_database (
                                contract_id, date_of_outflow, name_surname, amount, status
                            ) VALUES (%s, %s, %s, %s, %s)
                        """, (
                contract_id,
                self.take_date_box.text(),
                self.name_surname_box.text(),
                self.given_money_box.text(),
                status_for_given_principle
            ))

            new_datetime_str = QDateTime.fromString(self.take_date_box.text(), "yyyy-MM-dd HH:mm:ss") \
                .addDays(int(self.day_quantity_box.currentText()) - 1) \
                .toString("yyyy-MM-dd HH:mm:ss")

            cursor.execute("""
                            INSERT INTO contracts (
                                contract_id, contract_open_date, first_percent_payment_date, name_surname, id_number, 
                                tel_number, item_name, model, IMEI, given_money,percent_day_quantity, 
                                first_added_percent, office_mob_number, comment, trusted_person
                            ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                        """, (
                contract_id,
                self.take_date_box.text(),
                new_datetime_str,
                self.name_surname_box.text(),
                self.id_number_box.text(),
                self.tel_number_box.text(),
                self.item_name_box.text(),
                self.model_box.text(),
                self.imei_sn_box.text(),
                self.given_money_box.text(),
                int(self.day_quantity_box.currentText()),
                float(int(self.given_money_box.text()) * float(self.percent_box.currentText()) / 100),
                office_mob_number,
                self.comment_box.text(),
                self.trusted_person_box.text()
            ))

            cursor.execute("""
                                       INSERT INTO outflow_order (
                                           contract_id, date, name_surname, tel_number, amount, status
                                       ) VALUES (%s, %s, %s, %s, %s, %s)
                                   """, (
                contract_id,
                self.take_date_box.text(),
                self.name_surname_box.text(),
                self.tel_number_box.text(),
                self.given_money_box.text(),
                status_for_given_principle
            ))

            status_for_added_percent = "დარიცხული პროცენტი"


            cursor.execute("""
                            INSERT INTO adding_percent_amount (
                                    contract_id, date_of_C_O, name_surname, id_number, 
                                    tel_number, item_name, model, IMEI, date_of_percent_addition, percent_amount, status
                                    ) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
                            """, (
                contract_id,
                self.take_date_box.text(),
                self.name_surname_box.text(),
                self.id_number_box.text(),
                self.tel_number_box.text(),
                self.item_name_box.text(),
                self.model_box.text(),
                self.imei_sn_box.text(),
                self.take_date_box.text(),
                int(self.given_money_box.text()) * float(self.percent_box.currentText()) / 100,
                status_for_added_percent
            ))

            conn.commit()
            conn.close()

            QMessageBox.information(self, "წარმატება", "მონაცემები შენახულია")
            self.close()


            # Printing opening contract
            dt = datetime.strptime(self.take_date_box.text(), "%Y-%m-%d %H:%M:%S")
            date = dt.strftime("%d-%m-%Y")

            replacements = {
                '{name_surname}': self.name_surname_box.text() or "",
                '{given_money}': str(self.given_money_box.text()) if self.given_money_box.text() is not None else "",
                '{date}': date or "",
                '{contract_id}': str(contract_id or ""),
                '{id_number}': self.id_number_box.text() or "",
                '{IMEI}': self.imei_sn_box.text() or "",
                '{model}': self.model_box.text() or "",
                '{item_name}': self.item_name_box.text() or "",
                '{comment}': self.comment_box.text() or "",
                '{trusted_person}': self.trusted_person_box.text() or "",
                '{tel_number}': self.tel_number_box.text() or "",
                '{organization_name}': getattr(self, "organisation", ""),
                '{operator_name}': getattr(self, "name_of_user", "")
            }

            def replace_in_paragraph(paragraph, replacements):
                full_text = ''.join(run.text for run in paragraph.runs)
                new_text = full_text
                for key, value in replacements.items():
                    new_text = new_text.replace(key, str(value))

                if new_text != full_text:
                    for run in paragraph.runs:
                        run.text = ''
                    if paragraph.runs:
                        paragraph.runs[0].text = new_text
                    else:
                        paragraph.add_run(new_text)

            # Load the Word template
            doc = Document(resource_path("Templates/contract_template.docx"))

            # Replace in normal paragraphs
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph, replacements)

            # Replace in table cells
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph, replacements)

            output_dir = "GeneratedContracts"
            os.makedirs(output_dir, exist_ok=True)

            # Construct file name
            output_filename = f"contract_{contract_id}_{self.name_surname_box.text()}.docx"
            output_path = os.path.join(output_dir, output_filename)

            # Save document
            doc.save(output_path)

            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word_doc = word.Documents.Open(os.path.abspath(output_path))
                word_doc.PrintOut()  # Try to print
                word_doc.Close(False)  # Close document without saving
                word.Quit()
            except Exception as e:
                print("Printing failed:", e)
                print("Opening document instead...")
                try:
                    os.startfile(output_path)  # Open in Word as fallback
                except Exception as open_error:
                    print("Could not open Word document:", open_error)



        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", f"ვერ შევინახე მონაცემები:\n{e}")
        finally:
            conn.close()
