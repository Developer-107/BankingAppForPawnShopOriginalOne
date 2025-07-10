import os
import sqlite3
import tempfile
from datetime import datetime
import win32com.client
import pandas as pd
from PyQt5.QtCore import QDate, QSize, Qt, QPoint
from PyQt5.QtGui import QIcon, QTextDocument
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
from PyQt5.QtSql import QSqlDatabase, QSqlTableModel
from PyQt5.QtWidgets import QWidget, QGridLayout, QPushButton, QHBoxLayout, QTabWidget, QLabel, QVBoxLayout, QGroupBox, \
    QDateEdit, QTableView, QAbstractItemView, QToolButton, QRadioButton, QLineEdit, QButtonGroup, QMessageBox, QMenu, \
    QAction
from docx import Document
from adding_percent_amount_edit_window import EditAddingPercentWindow
from outflow_in_registry_edit_money_control_window import EditRegistryOutflowWindow
from open_edit_money_control_window_2 import EditInPrincipalInflowsInRegistryWindow
from registry_page_4_edit_window_4_paid_percents import EditPaidPercentWindow


class ContractRegistry(QWidget):
    def __init__(self, role, name_of_user, organisation):
        super().__init__()
        self.setWindowTitle("ხელშეკრულებების რეესტრი")
        self.setWindowIcon(QIcon("Icons/contract_registry.png"))
        self.resize(1701, 700)
        self.role = role
        self.organisation = organisation
        self.name_of_user = name_of_user

        layout = QGridLayout()

        self.tabs = QTabWidget()
        self.tabs.setStyleSheet("""
            QTabBar::tab {
                font-size: 8pt;
                font-weight: bold;
                padding: 6px 16px;
                background: #f0f0f0;
                border: 1px solid #ccc;
                border-bottom: none;
                min-width: 197px;
                max-width: 197px;
                min-height: 25px;
                border-top-left-radius: 5px;
                border-top-right-radius: 5px;
            }
            QTabBar::tab:selected {
                background: white;
                color: black;
            }
            QTabWidget::pane {
                border: 1px solid #ccc;
                top: -1px;
            }
        """)


        # ---------------------------------------------Page1--------------------------------------------------------

        self.page1 = QWidget()
        layout1 = QGridLayout()

        box1 = QGroupBox("ძებნა")
        box1.setFixedWidth(400)
        box1.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout1 = QGridLayout()

        self.from_date = QDateEdit()
        self.from_date.setCalendarPopup(True)
        self.from_date.setDate(QDate.currentDate().addMonths(-1))
        self.to_date = QDateEdit()
        self.to_date.setCalendarPopup(True)
        self.to_date.setDate(QDate.currentDate())

        self.contract_date_radio = QRadioButton("ხელშეკრულების გაფორმების თარიღით")
        self.closing_date_radio = QRadioButton("დამატების თარიღით")
        self.closing_date_radio.setChecked(True)  # Default


        box_layout1.addWidget(self.contract_date_radio, 0, 0, 1, 2)
        box_layout1.addWidget(self.closing_date_radio, 1, 0, 1, 2)

        box_layout1.addWidget(QLabel("*დან თარიღი:"), 2, 0)
        box_layout1.addWidget(self.from_date, 2, 1)
        box_layout1.addWidget(QLabel("*მდე თარიღი:"), 3, 0)
        box_layout1.addWidget(self.to_date, 3, 1)


        filter_button = QPushButton("ძებნა")
        filter_button.clicked.connect(self.search_by_date)
        refresh_button = QPushButton("განახლება")
        refresh_button.clicked.connect(self.refresh_table)

        box_layout1.addWidget(refresh_button)
        box_layout1.addWidget(filter_button)

        layout1.addWidget(box1, 0, 2, 1, 2)
        box1.setLayout(box_layout1)

        box2 = QGroupBox("ძებნა")
        box2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout2 = QGridLayout()

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("მოძებნე აქ")
        self.search_input.textChanged.connect(self.apply_text_filter)

        # Radio Button
        self.contract_radio = QRadioButton("ხელშეკრულების N")
        self.name_radio = QRadioButton("სახელი და გვარი")
        self.id_radio = QRadioButton("პირადი N")
        self.model_radio = QRadioButton("მოდელი")
        self.tel_radio = QRadioButton("ტელეფონის ნომერი")
        self.imei_radio = QRadioButton("IMEI")

        # Grouping buttons
        button_group = QButtonGroup()
        button_group.addButton(self.contract_radio)
        button_group.addButton(self.name_radio)
        button_group.addButton(self.id_radio)
        button_group.addButton(self.model_radio)
        button_group.addButton(self.tel_radio)
        button_group.addButton(self.imei_radio)

        # Making Widgets
        box_layout2.addWidget(self.contract_radio)
        box_layout2.addWidget(self.name_radio)
        box_layout2.addWidget(self.id_radio)
        box_layout2.addWidget(self.model_radio)
        box_layout2.addWidget(self.tel_radio)
        box_layout2.addWidget(self.imei_radio)
        box_layout2.addWidget(self.search_input, 2, 1)

        layout1.addWidget(box2, 0, 1, 1, 1)
        box2.setLayout(box_layout2)

        box3 = QGroupBox("ნავიგაცია")
        box3.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout3 = QGridLayout()

        # Box 3 buttons
        # Edit
        edit_money_control = QToolButton()
        edit_money_control.setText(" რედაქტირება ")
        # edit_money_control.setIcon(QIcon("Icons/.png"))
        edit_money_control.setIconSize(QSize(37, 40))
        edit_money_control.setFixedHeight(45)
        edit_money_control.setFixedWidth(145)
        edit_money_control.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        edit_money_control.setStyleSheet("font-size: 16px;")
        edit_money_control.clicked.connect(self.open_outflow_in_registry_edit_money_control_window)
        if self.role != "admin":
            edit_money_control.setEnabled(False)

        # Remove
        delete_from_money_control = QToolButton()
        delete_from_money_control.setText(" ამოშლა ")
        # delete_from_money_control.setIcon(QIcon("Icons/.png"))
        delete_from_money_control.setIconSize(QSize(37, 40))
        delete_from_money_control.setFixedHeight(45)
        delete_from_money_control.setFixedWidth(145)
        delete_from_money_control.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        delete_from_money_control.setStyleSheet("font-size: 16px;")
        delete_from_money_control.clicked.connect(self.delete_selected_row)
        if self.role != "admin":
            delete_from_money_control.setEnabled(False)

        # Export
        export_money_control = QToolButton()
        export_money_control.setText(" ექსპორტი ")
        export_money_control.setIcon(QIcon("Icons/excel_icon.png"))
        export_money_control.setIconSize(QSize(37, 40))
        export_money_control.setFixedHeight(45)
        export_money_control.setFixedWidth(145)
        export_money_control.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        export_money_control.setStyleSheet("font-size: 16px;")
        export_money_control.clicked.connect(self.export_to_excel)

        # Box3 widgets
        box_layout3.addWidget(edit_money_control, 1, 0)
        box_layout3.addWidget(delete_from_money_control, 2, 0)
        box_layout3.addWidget(export_money_control, 3, 0)

        layout1.addWidget(box3, 1, 0)
        box3.setLayout(box_layout3)



        self.db2 = QSqlDatabase.addDatabase("QSQLITE", "outflow_in_registry")
        self.db2.setDatabaseName("Databases/outflow_in_registry.db")
        if not self.db2.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table2 = QTableView()
        self.model2 = QSqlTableModel(self, self.db2)
        self.model2.setTable("outflow_in_registry")
        self.model2.select()
        self.table2.setModel(self.model2)
        self.table2.setEditTriggers(QAbstractItemView.NoEditTriggers)  # Read-only table
        self.table2.setSelectionBehavior(QTableView.SelectRows)
        self.table2.setSelectionMode(QTableView.SingleSelection)
        self.table2.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table2.customContextMenuRequested.connect(self.show_registry_page1_table)

        layout1.addWidget(self.table2, 1, 1, 2, 3)





        # Page1 Layout
        self.page1.setLayout(layout1)

        # --------------------------------------------Page2--------------------------------------------------------

        self.page2 = QWidget()
        layout2 = QGridLayout()

        box2_1 = QGroupBox("ძებნა")
        box2_1.setFixedWidth(400)
        box2_1.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout2_1 = QGridLayout()

        self.from_date_2 = QDateEdit()
        self.from_date_2.setCalendarPopup(True)
        self.from_date_2.setDate(QDate.currentDate().addMonths(-1))
        self.to_date_2 = QDateEdit()
        self.to_date_2.setCalendarPopup(True)
        self.to_date_2.setDate(QDate.currentDate())

        self.date_radio_2 = QRadioButton("ხელშეკრულების გაფორმების თარიღით")
        self.closing_date_radio_2 = QRadioButton("ძირის შემოტანის თარიღით")
        self.closing_date_radio_2.setChecked(True)

        box_layout2_1.addWidget(self.date_radio_2, 0, 0, 1, 2)
        box_layout2_1.addWidget(self.closing_date_radio_2, 1, 0, 1, 2)
        box_layout2_1.addWidget(QLabel("*დან თარიღი:"), 2, 0)
        box_layout2_1.addWidget(self.from_date_2, 2, 1)
        box_layout2_1.addWidget(QLabel("*მდე თარიღი:"), 3, 0)
        box_layout2_1.addWidget(self.to_date_2, 3, 1)

        search_button_2 = QPushButton("ძებნა")
        search_button_2.clicked.connect(self.search_by_date_2)
        refresh_button_2 = QPushButton("განახლება")
        refresh_button_2.clicked.connect(self.refresh_table_2)

        box_layout2_1.addWidget(refresh_button_2)
        box_layout2_1.addWidget(search_button_2)

        box2_1.setLayout(box_layout2_1)
        layout2.addWidget(box2_1, 0, 2, 1, 2)

        box2_2 = QGroupBox("ძებნა")
        box2_2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout2_2 = QGridLayout()

        self.search_input_2 = QLineEdit()
        self.search_input_2.setPlaceholderText("მოძებნე აქ")
        self.search_input_2.textChanged.connect(self.apply_text_filter_2)

        # Radio buttons
        self.contract_radio_2 = QRadioButton("ხელშეკრულების N")
        self.name_radio_2 = QRadioButton("სახელი და გვარი")
        self.id_radio_2 = QRadioButton("პირადი N")
        self.model_radio_2 = QRadioButton("მოდელი")
        self.tel_radio_2 = QRadioButton("ტელეფონის ნომერი")
        self.imei_radio_2 = QRadioButton("IMEI")

        # Group radio buttons
        search_group_2 = QButtonGroup()
        search_group_2.addButton(self.contract_radio_2)
        search_group_2.addButton(self.name_radio_2)
        search_group_2.addButton(self.id_radio_2)
        search_group_2.addButton(self.model_radio_2)
        search_group_2.addButton(self.tel_radio_2)
        search_group_2.addButton(self.imei_radio_2)

        # Add to layout (organized)
        box_layout2_2.addWidget(self.contract_radio_2, 0, 0)
        box_layout2_2.addWidget(self.name_radio_2, 1, 0)
        box_layout2_2.addWidget(self.id_radio_2, 2, 0)
        box_layout2_2.addWidget(self.model_radio_2, 3, 0)
        box_layout2_2.addWidget(self.tel_radio_2, 4, 0)
        box_layout2_2.addWidget(self.imei_radio_2, 5, 0)
        box_layout2_2.addWidget(self.search_input_2, 2, 1)

        box2_2.setLayout(box_layout2_2)
        layout2.addWidget(box2_2, 0, 1, 1, 1)

        # --- Box 3: Navigation Buttons ---
        box2_3 = QGroupBox("ნავიგაცია")
        box2_3.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout2_3 = QGridLayout()

        edit_button_2 = QToolButton()
        edit_button_2.setText(" რედაქტირება ")
        edit_button_2.setFixedSize(145, 45)
        edit_button_2.setStyleSheet("font-size: 16px;")
        edit_button_2.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        edit_button_2.clicked.connect(self.open_edit_money_control_window_2)
        if self.role != "admin":
            edit_button_2.setEnabled(False)

        delete_button_2 = QToolButton()
        delete_button_2.setText(" ამოშლა ")
        delete_button_2.setFixedSize(145, 45)
        delete_button_2.setStyleSheet("font-size: 16px;")
        delete_button_2.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        delete_button_2.clicked.connect(self.delete_selected_row_2)
        if self.role != "admin":
            delete_button_2.setEnabled(False)

        export_button_2 = QToolButton()
        export_button_2.setText(" ექსპორტი ")
        export_button_2.setIcon(QIcon("Icons/excel_icon.png"))
        export_button_2.setIconSize(QSize(37, 40))
        export_button_2.setFixedSize(145, 45)
        export_button_2.setStyleSheet("font-size: 16px;")
        export_button_2.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        export_button_2.clicked.connect(self.export_to_excel_2)

        box_layout2_3.addWidget(edit_button_2, 1, 0)
        box_layout2_3.addWidget(delete_button_2, 2, 0)
        box_layout2_3.addWidget(export_button_2, 3, 0)

        box2_3.setLayout(box_layout2_3)
        layout2.addWidget(box2_3, 1, 0)

        # --- Table Setup ---
        self.db3 = QSqlDatabase.addDatabase("QSQLITE","paid_principle_registry")
        self.db3.setDatabaseName("Databases/paid_principle_registry.db")
        if not self.db3.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table3 = QTableView()
        self.model3 = QSqlTableModel(self, self.db3)
        self.model3.setTable("paid_principle_registry")  # or another table name
        self.model3.select()
        self.table3.setModel(self.model3)
        self.table3.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table3.setSelectionBehavior(QTableView.SelectRows)
        self.table3.setSelectionMode(QTableView.SingleSelection)

        layout2.addWidget(self.table3, 1, 1, 2, 3)


        # Page2 layout
        self.page2.setLayout(layout2)

        # --------------------------------------------Page3--------------------------------------------------------

        self.page3 = QWidget()
        layout3 = QGridLayout()

        # --- Box 1: Date Filter ---
        box3_1 = QGroupBox("ძებნა")
        box3_1.setFixedWidth(400)
        box3_1.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout3_1 = QGridLayout()

        self.from_date_3 = QDateEdit()
        self.from_date_3.setCalendarPopup(True)
        self.from_date_3.setDate(QDate.currentDate().addMonths(-1))
        self.to_date_3 = QDateEdit()
        self.to_date_3.setCalendarPopup(True)
        self.to_date_3.setDate(QDate.currentDate())

        self.date_radio_3 = QRadioButton("ხელშეკრულების გაფორმების თარიღით")
        self.closing_date_radio_3 = QRadioButton("პროცენტის დარიცხვის თარიღით")
        self.closing_date_radio_3.setChecked(True)

        box_layout3_1.addWidget(self.date_radio_3, 0, 0, 1, 2)
        box_layout3_1.addWidget(self.closing_date_radio_3, 1, 0, 1, 2)
        box_layout3_1.addWidget(QLabel("*დან თარიღი:"), 2, 0)
        box_layout3_1.addWidget(self.from_date_3, 2, 1)
        box_layout3_1.addWidget(QLabel("*მდე თარიღი:"), 3, 0)
        box_layout3_1.addWidget(self.to_date_3, 3, 1)

        search_button_3 = QPushButton("ძებნა")
        search_button_3.clicked.connect(self.search_by_date_3)
        refresh_button_3 = QPushButton("განახლება")
        refresh_button_3.clicked.connect(self.refresh_table_3)

        box_layout3_1.addWidget(refresh_button_3)
        box_layout3_1.addWidget(search_button_3)

        box3_1.setLayout(box_layout3_1)
        layout3.addWidget(box3_1, 0, 2, 1, 2)

        # --- Box 2: Text Filter ---
        box3_2 = QGroupBox("ძებნა")
        box3_2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout3_2 = QGridLayout()

        self.search_input_3 = QLineEdit()
        self.search_input_3.setPlaceholderText("მოძებნე აქ")
        self.search_input_3.textChanged.connect(self.apply_text_filter_3)

        self.contract_radio_3 = QRadioButton("ხელშეკრულების N")
        self.name_radio_3 = QRadioButton("სახელი და გვარი")
        self.id_radio_3 = QRadioButton("პირადი N")
        self.model_radio_3 = QRadioButton("მოდელი")
        self.tel_radio_3 = QRadioButton("ტელეფონის ნომერი")
        self.imei_radio_3 = QRadioButton("IMEI")

        search_group_3 = QButtonGroup()
        search_group_3.addButton(self.contract_radio_3)
        search_group_3.addButton(self.name_radio_3)
        search_group_3.addButton(self.id_radio_3)
        search_group_3.addButton(self.model_radio_3)
        search_group_3.addButton(self.tel_radio_3)
        search_group_3.addButton(self.imei_radio_3)

        box_layout3_2.addWidget(self.contract_radio_3, 0, 0)
        box_layout3_2.addWidget(self.name_radio_3, 1, 0)
        box_layout3_2.addWidget(self.id_radio_3, 2, 0)
        box_layout3_2.addWidget(self.model_radio_3, 3, 0)
        box_layout3_2.addWidget(self.tel_radio_3, 4, 0)
        box_layout3_2.addWidget(self.imei_radio_3, 5, 0)
        box_layout3_2.addWidget(self.search_input_3, 2, 1)

        box3_2.setLayout(box_layout3_2)
        layout3.addWidget(box3_2, 0, 1, 1, 1)

        # --- Box 3: Navigation Buttons ---
        box3_3 = QGroupBox("ნავიგაცია")
        box3_3.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout3_3 = QGridLayout()

        edit_button_3 = QToolButton()
        edit_button_3.setText(" რედაქტირება ")
        edit_button_3.setFixedSize(145, 45)
        edit_button_3.setStyleSheet("font-size: 16px;")
        edit_button_3.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        edit_button_3.clicked.connect(self.open_edit_money_control_window_3)
        if self.role != "admin":
            edit_button_3.setEnabled(False)

        delete_button_3 = QToolButton()
        delete_button_3.setText(" ამოშლა ")
        delete_button_3.setFixedSize(145, 45)
        delete_button_3.setStyleSheet("font-size: 16px;")
        delete_button_3.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        delete_button_3.clicked.connect(self.delete_selected_row_3)
        if self.role != "admin":
            delete_button_3.setEnabled(False)

        export_button_3 = QToolButton()
        export_button_3.setText(" ექსპორტი ")
        export_button_3.setIcon(QIcon("Icons/excel_icon.png"))
        export_button_3.setIconSize(QSize(37, 40))
        export_button_3.setFixedSize(145, 45)
        export_button_3.setStyleSheet("font-size: 16px;")
        export_button_3.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        export_button_3.clicked.connect(self.export_to_excel_3)

        box_layout3_3.addWidget(edit_button_3, 1, 0)
        box_layout3_3.addWidget(delete_button_3, 2, 0)
        box_layout3_3.addWidget(export_button_3, 3, 0)

        box3_3.setLayout(box_layout3_3)
        layout3.addWidget(box3_3, 0, 0)

        # --- Table Setup ---
        self.db3_1 = QSqlDatabase.addDatabase("QSQLITE", "adding_percent_amount")
        self.db3_1.setDatabaseName("Databases/adding_percent_amount.db")
        if not self.db3_1.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table3_1 = QTableView()
        self.model3_1 = QSqlTableModel(self, self.db3_1)
        self.model3_1.setTable("adding_percent_amount")
        self.model3_1.select()
        self.table3_1.setModel(self.model3_1)
        self.table3_1.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table3_1.setSelectionBehavior(QTableView.SelectRows)
        self.table3_1.setSelectionMode(QTableView.SingleSelection)

        layout3.addWidget(self.table3_1, 1, 0, 2, 4)

        # Set layout
        self.page3.setLayout(layout3)

        # --------------------------------------------Page4--------------------------------------------------------

        self.page4 = QWidget()
        layout4 = QGridLayout()

        # --- Box 1: Date Filter ---
        box4_1 = QGroupBox("ძებნა")
        box4_1.setFixedWidth(400)
        box4_1.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout4_1 = QGridLayout()

        self.from_date_4 = QDateEdit()
        self.from_date_4.setCalendarPopup(True)
        self.from_date_4.setDate(QDate.currentDate().addMonths(-1))
        self.to_date_4 = QDateEdit()
        self.to_date_4.setCalendarPopup(True)
        self.to_date_4.setDate(QDate.currentDate())

        self.date_radio_4 = QRadioButton("ხელშეკრულების გაფორმების თარიღით")
        self.closing_date_radio_4 = QRadioButton("პროცენტის გადახდის თარიღით")
        self.closing_date_radio_4.setChecked(True)

        box_layout4_1.addWidget(self.date_radio_4, 0, 0, 1, 2)
        box_layout4_1.addWidget(self.closing_date_radio_4, 1, 0, 1, 2)
        box_layout4_1.addWidget(QLabel("*დან თარიღი:"), 2, 0)
        box_layout4_1.addWidget(self.from_date_4, 2, 1)
        box_layout4_1.addWidget(QLabel("*მდე თარიღი:"), 3, 0)
        box_layout4_1.addWidget(self.to_date_4, 3, 1)

        search_button_4 = QPushButton("ძებნა")
        search_button_4.clicked.connect(self.search_by_date_4)
        refresh_button_4 = QPushButton("განახლება")
        refresh_button_4.clicked.connect(self.refresh_table_4)

        box_layout4_1.addWidget(refresh_button_4)
        box_layout4_1.addWidget(search_button_4)

        box4_1.setLayout(box_layout4_1)
        layout4.addWidget(box4_1, 0, 2, 1, 2)

        # --- Box 2: Text Filter ---
        box4_2 = QGroupBox("ძებნა")
        box4_2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout4_2 = QGridLayout()

        self.search_input_4 = QLineEdit()
        self.search_input_4.setPlaceholderText("მოძებნე აქ")
        self.search_input_4.textChanged.connect(self.apply_text_filter_4)

        self.contract_radio_4 = QRadioButton("ხელშეკრულების N")
        self.name_radio_4 = QRadioButton("სახელი და გვარი")
        self.id_radio_4 = QRadioButton("პირადი N")
        self.model_radio_4 = QRadioButton("მოდელი")
        self.tel_radio_4 = QRadioButton("ტელეფონის ნომერი")
        self.imei_radio_4 = QRadioButton("IMEI")

        search_group_4 = QButtonGroup()
        search_group_4.addButton(self.contract_radio_4)
        search_group_4.addButton(self.name_radio_4)
        search_group_4.addButton(self.id_radio_4)
        search_group_4.addButton(self.model_radio_4)
        search_group_4.addButton(self.tel_radio_4)
        search_group_4.addButton(self.imei_radio_4)

        box_layout4_2.addWidget(self.contract_radio_4, 0, 0)
        box_layout4_2.addWidget(self.name_radio_4, 1, 0)
        box_layout4_2.addWidget(self.id_radio_4, 2, 0)
        box_layout4_2.addWidget(self.model_radio_4, 3, 0)
        box_layout4_2.addWidget(self.tel_radio_4, 4, 0)
        box_layout4_2.addWidget(self.imei_radio_4, 5, 0)
        box_layout4_2.addWidget(self.search_input_4, 2, 1)

        box4_2.setLayout(box_layout4_2)
        layout4.addWidget(box4_2, 0, 1, 1, 1)

        # --- Box 3: Navigation Buttons ---
        box4_3 = QGroupBox("ნავიგაცია")
        box4_3.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout4_3 = QGridLayout()

        edit_button_4 = QToolButton()
        edit_button_4.setText(" რედაქტირება ")
        edit_button_4.setFixedSize(145, 45)
        edit_button_4.setStyleSheet("font-size: 16px;")
        edit_button_4.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        edit_button_4.clicked.connect(self.open_edit_money_control_window_4)
        if self.role != "admin":
            edit_button_4.setEnabled(False)


        delete_button_4 = QToolButton()
        delete_button_4.setText(" ამოშლა ")
        delete_button_4.setFixedSize(145, 45)
        delete_button_4.setStyleSheet("font-size: 16px;")
        delete_button_4.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        delete_button_4.clicked.connect(self.delete_selected_row_4)
        if self.role != "admin":
            delete_button_4.setEnabled(False)

        export_button_4 = QToolButton()
        export_button_4.setText(" ექსპორტი ")
        export_button_4.setIcon(QIcon("Icons/excel_icon.png"))
        export_button_4.setIconSize(QSize(37, 40))
        export_button_4.setFixedSize(145, 45)
        export_button_4.setStyleSheet("font-size: 16px;")
        export_button_4.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        export_button_4.clicked.connect(self.export_to_excel_4)

        box_layout4_3.addWidget(edit_button_4, 1, 0)
        box_layout4_3.addWidget(delete_button_4, 2, 0)
        box_layout4_3.addWidget(export_button_4, 3, 0)

        box4_3.setLayout(box_layout4_3)
        layout4.addWidget(box4_3, 1, 0)

        # --- Table Setup ---
        self.db4 = QSqlDatabase.addDatabase("QSQLITE", "paid_percent_amount")
        self.db4.setDatabaseName("Databases/paid_percent_amount.db")
        if not self.db4.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table4 = QTableView()
        self.model4 = QSqlTableModel(self, self.db4)
        self.model4.setTable("paid_percent_amount")
        self.model4.select()
        self.table4.setModel(self.model4)
        self.table4.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table4.setSelectionBehavior(QTableView.SelectRows)
        self.table4.setSelectionMode(QTableView.SingleSelection)

        layout4.addWidget(self.table4, 1, 1, 2, 3)

        # Set final layout
        self.page4.setLayout(layout4)

        # --------------------------------------------Page5--------------------------------------------------------

        self.page5 = QWidget()
        layout5 = QGridLayout()

        # --- Box 1: Date Filter ---
        box5_1 = QGroupBox("ძებნა")
        box5_1.setFixedWidth(400)
        box5_1.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout5_1 = QGridLayout()



        # --- Box 2: Text Filter ---
        box5_2 = QGroupBox("ძებნა")
        box5_2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout5_2 = QGridLayout()

        self.search_input_5 = QLineEdit()
        self.search_input_5.setPlaceholderText("მოძებნე აქ")
        self.search_input_5.textChanged.connect(self.apply_text_filter_5)


        self.contract_radio_5 = QRadioButton("ხელშეკრულების N")
        self.name_radio_5 = QRadioButton("სახელი და გვარი")
        self.tel_radio_5 = QRadioButton("ტელეფონის ნომერი")

        search_group_5 = QButtonGroup()
        search_group_5.addButton(self.contract_radio_5)
        search_group_5.addButton(self.name_radio_5)
        search_group_5.addButton(self.tel_radio_5)

        box_layout5_2.addWidget(self.contract_radio_5, 0, 0)
        box_layout5_2.addWidget(self.name_radio_5, 0, 1)
        box_layout5_2.addWidget(self.tel_radio_5, 0, 2)
        box_layout5_2.addWidget(self.search_input_5, 1, 0, 1, 3)

        box5_2.setLayout(box_layout5_2)
        layout5.addWidget(box5_2, 0, 0, 1, 2)


        # --- Table Setup ---
        self.db5 = QSqlDatabase.addDatabase("QSQLITE", "outflow_order")
        self.db5.setDatabaseName("Databases/outflow_order.db")
        if not self.db5.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხქერდა")

        self.table5 = QTableView()
        self.model5 = QSqlTableModel(self, self.db5)
        self.model5.setTable("outflow_order")
        self.model5.select()
        self.table5.setModel(self.model5)
        self.table5.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table5.setSelectionBehavior(QTableView.SelectRows)
        self.table5.setSelectionMode(QTableView.SingleSelection)
        self.table5.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table5.customContextMenuRequested.connect(self.show_table5_context_menu)

        layout5.addWidget(self.table5, 1, 0, 4, 4)

        # Set layout
        self.page5.setLayout(layout5)




        # --------------------------------------------Page6--------------------------------------------------------

        self.page6 = QWidget()
        layout6 = QVBoxLayout(self.page6)

        self.sub_tab_widget_6 = QTabWidget()
        layout6.addWidget(self.sub_tab_widget_6)

        self.tab6_1 = QWidget()
        self.tab6_2 = QWidget()
        self.tab6_3 = QWidget()

        self.tab6_1 = QWidget()
        layout_tab6_1 = QGridLayout()

        # --- Box 2: Text Filter ---
        box6_1_2 = QGroupBox("ძებნა")
        box6_1_2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout6_1_2 = QGridLayout()

        self.search_input_6_1 = QLineEdit()
        self.search_input_6_1.setPlaceholderText("მოძებნე აქ")
        self.search_input_6_1.textChanged.connect(self.apply_text_filter_6_1)

        self.contract_radio_6_1 = QRadioButton("ხელშეკრულების N")
        self.name_radio_6_1 = QRadioButton("სახელი და გვარი")

        search_group_6_1 = QButtonGroup()
        search_group_6_1.addButton(self.contract_radio_6_1)
        search_group_6_1.addButton(self.name_radio_6_1)

        box_layout6_1_2.addWidget(self.contract_radio_6_1, 0, 0)
        box_layout6_1_2.addWidget(self.name_radio_6_1, 0, 1)
        box_layout6_1_2.addWidget(self.search_input_6_1, 1, 0, 1, 2)

        box6_1_2.setLayout(box_layout6_1_2)
        layout_tab6_1.addWidget(box6_1_2, 0, 0, 1, 1)

        # --- Table Setup ---
        self.db6_1 = QSqlDatabase.addDatabase("QSQLITE", "inflow_order_only_percent_amount")
        self.db6_1.setDatabaseName("Databases/inflow_order_only_percent_amount.db")
        if not self.db6_1.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table6_1 = QTableView()
        self.model6_1 = QSqlTableModel(self, self.db6_1)
        self.model6_1.setTable("inflow_order_only_percent_amount")
        self.model6_1.select()
        self.table6_1.setModel(self.model6_1)
        self.table6_1.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table6_1.setSelectionBehavior(QTableView.SelectRows)
        self.table6_1.setSelectionMode(QTableView.SingleSelection)
        self.table6_1.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table6_1.customContextMenuRequested.connect(self.show_table6_1_context_menu)

        layout_tab6_1.addWidget(self.table6_1, 1, 0, 4, 4)

        # Set layout
        self.tab6_1.setLayout(layout_tab6_1)

        self.tab6_2 = QWidget()
        layout_tab6_2 = QGridLayout()

        # --- Box: Text Filter ---
        box6_2_2 = QGroupBox("ძებნა")
        box6_2_2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout6_2_2 = QGridLayout()

        self.search_input_6_2 = QLineEdit()
        self.search_input_6_2.setPlaceholderText("მოძებნე აქ")
        self.search_input_6_2.textChanged.connect(self.apply_text_filter_6_2)

        self.contract_radio_6_2 = QRadioButton("ხელშეკრულების N")
        self.name_radio_6_2 = QRadioButton("სახელი და გვარი")

        search_group_6_2 = QButtonGroup()
        search_group_6_2.addButton(self.contract_radio_6_2)
        search_group_6_2.addButton(self.name_radio_6_2)

        box_layout6_2_2.addWidget(self.contract_radio_6_2, 0, 0)
        box_layout6_2_2.addWidget(self.name_radio_6_2, 0, 1)
        box_layout6_2_2.addWidget(self.search_input_6_2, 1, 0, 1, 2)

        box6_2_2.setLayout(box_layout6_2_2)
        layout_tab6_2.addWidget(box6_2_2, 0, 0, 1, 1)

        # --- Table Setup ---
        self.db6_2 = QSqlDatabase.addDatabase("QSQLITE", "inflow_order_only_principal_amount")
        self.db6_2.setDatabaseName("Databases/inflow_order_only_principal_amount.db")
        if not self.db6_2.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table6_2 = QTableView()
        self.model6_2 = QSqlTableModel(self, self.db6_2)
        self.model6_2.setTable("inflow_order_only_principal_amount")
        self.model6_2.select()
        self.table6_2.setModel(self.model6_2)
        self.table6_2.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table6_2.setSelectionBehavior(QTableView.SelectRows)
        self.table6_2.setSelectionMode(QTableView.SingleSelection)
        self.table6_2.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table6_2.customContextMenuRequested.connect(self.show_table6_2_context_menu)

        layout_tab6_2.addWidget(self.table6_2, 1, 0, 4, 4)

        self.tab6_2.setLayout(layout_tab6_2)



        self.tab6_3 = QWidget()
        layout_tab6_3 = QGridLayout()

        # --- Box: Text Filter ---
        box6_3_2 = QGroupBox("ძებნა")
        box6_3_2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout6_3_2 = QGridLayout()

        self.search_input_6_3 = QLineEdit()
        self.search_input_6_3.setPlaceholderText("მოძებნე აქ")
        self.search_input_6_3.textChanged.connect(self.apply_text_filter_6_3)


        self.contract_radio_6_3 = QRadioButton("ხელშეკრულების N")
        self.name_radio_6_3 = QRadioButton("სახელი და გვარი")

        search_group_6_3 = QButtonGroup()
        search_group_6_3.addButton(self.contract_radio_6_3)
        search_group_6_3.addButton(self.name_radio_6_3)

        box_layout6_3_2.addWidget(self.contract_radio_6_3, 0, 0)
        box_layout6_3_2.addWidget(self.name_radio_6_3, 0, 1)
        box_layout6_3_2.addWidget(self.search_input_6_3, 1, 0, 1, 2)

        box6_3_2.setLayout(box_layout6_3_2)
        layout_tab6_3.addWidget(box6_3_2, 0, 0, 1, 1)

        # --- Table Setup ---
        self.db6_3 = QSqlDatabase.addDatabase("QSQLITE", "inflow_order_both")
        self.db6_3.setDatabaseName("Databases/inflow_order_both.db")
        if not self.db6_3.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table6_3 = QTableView()
        self.model6_3 = QSqlTableModel(self, self.db6_3)
        self.model6_3.setTable("inflow_order_both")
        self.model6_3.select()
        self.table6_3.setModel(self.model6_3)
        self.table6_3.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table6_3.setSelectionBehavior(QTableView.SelectRows)
        self.table6_3.setSelectionMode(QTableView.SingleSelection)
        self.table6_3.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table6_3.customContextMenuRequested.connect(self.show_table6_3_context_menu)

        layout_tab6_3.addWidget(self.table6_3, 1, 0, 4, 4)

        self.tab6_3.setLayout(layout_tab6_3)



        self.sub_tab_widget_6.addTab(self.tab6_1, "მარტო პროცენტები")
        self.sub_tab_widget_6.addTab(self.tab6_2, "მარტო ძირი თანხები")
        self.sub_tab_widget_6.addTab(self.tab6_3, "ძირი თანხა და პროცენტები ერთად")

        self.sub_tab_widget_6.setStyleSheet("""
        QTabBar::tab {
            border: 2px solid #cccccc;
            border-bottom: none;
            min-width: 301px;
            border-top-left-radius: 5px;
            border-top-right-radius: 5px;
        }
        QTabBar::tab:selected {
            background: grey;
            color: white;
            border-color: black;
        }
        }
        QTabWidget::pane {
            border: 1px solid #999;
            top: -1px;
        }
        """)







        # --------------------------------------------Page7--------------------------------------------------------

        self.page7 = QWidget()
        layout7 = QGridLayout()

        # --- Box 2: Text Filter ---
        box7_2 = QGroupBox("ძებნა")
        box7_2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout7_2 = QGridLayout()

        self.search_input_7 = QLineEdit()
        self.search_input_7.setPlaceholderText("მოძებნე აქ")
        self.search_input_7.textChanged.connect(self.apply_text_filter_7)

        self.contract_radio_7 = QRadioButton("ხელშეკრულების N")
        self.name_radio_7 = QRadioButton("სახელი და გვარი")
        self.id_radio_7 = QRadioButton("პირადი N")
        self.model_radio_7 = QRadioButton("მოდელი")
        self.tel_radio_7 = QRadioButton("ტელეფონის ნომერი")
        self.imei_radio_7 = QRadioButton("IMEI")

        search_group_7 = QButtonGroup()
        search_group_7.addButton(self.contract_radio_7)
        search_group_7.addButton(self.name_radio_7)
        search_group_7.addButton(self.id_radio_7)
        search_group_7.addButton(self.model_radio_7)
        search_group_7.addButton(self.tel_radio_7)
        search_group_7.addButton(self.imei_radio_7)

        box_layout7_2.addWidget(self.contract_radio_7, 0, 0)
        box_layout7_2.addWidget(self.name_radio_7, 0, 1)
        box_layout7_2.addWidget(self.id_radio_7, 0, 2)
        box_layout7_2.addWidget(self.model_radio_7, 1, 2)
        box_layout7_2.addWidget(self.tel_radio_7, 1, 1)
        box_layout7_2.addWidget(self.imei_radio_7, 1, 0)
        box_layout7_2.addWidget(self.search_input_7, 1, 3)

        box7_2.setLayout(box_layout7_2)
        layout7.addWidget(box7_2, 0, 0, 1, 2)

        # --- Table Setup ---
        self.db7 = QSqlDatabase.addDatabase("QSQLITE", "contracts_connection")
        self.db7.setDatabaseName("Databases/contracts.db")
        if not self.db7.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table7 = QTableView()
        self.model7 = QSqlTableModel(self, self.db7)
        self.model7.setTable("contracts_view")
        self.model7.select()
        self.table7.setModel(self.model7)
        self.table7.setEditTriggers(QAbstractItemView.NoEditTriggers)
        self.table7.setSelectionBehavior(QTableView.SelectRows)
        self.table7.setSelectionMode(QTableView.SingleSelection)
        self.table7.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table7.customContextMenuRequested.connect(self.show_table7_context_menu)

        layout7.addWidget(self.table7, 1, 0, 4, 4)

        # Set layout
        self.page7.setLayout(layout7)

        # --------------------------------------------Tabs--------------------------------------------------------

        self.tabs.addTab(self.page1, "დამატებული თანხები")
        self.tabs.addTab(self.page2, "შემოტანილი ძირი თახნა")
        self.tabs.addTab(self.page3, "დარიცხული პროცენტები")
        self.tabs.addTab(self.page4, "გადახდილი პროცენტები")
        self.tabs.addTab(self.page5, "გასავლის ორდერები")
        self.tabs.addTab(self.page6, "შემოსავლის ორდერები")
        self.tabs.addTab(self.page7, "ხელშეკრულებები")



        layout.addWidget(self.tabs)








        # --------------------------------------------Layout-----------------------------------------------------
        self.setLayout(layout)


    # --------------------------------------------page1functions-----------------------------------------------------
    def show_registry_page1_table(self, position: QPoint):
        index = self.table2.indexAt(position)
        if not index.isValid():
            return

        menu = QMenu()

        print_action = QAction(" დამატებული თანხის დოკუმენტის ამობეჭდვა ", self)
        print_action.setIcon(QIcon("Icons/printer_icon.png"))
        print_action.triggered.connect(self.print_registry_page1_selected_row)

        menu.addAction(print_action)

        menu.exec_(self.table2.viewport().mapToGlobal(position))

    def print_registry_page1_selected_row(self):
        selected = self.table2.selectionModel().selectedRows()
        if not selected:
            print("No row selected.")
            return

        row_index = selected[0].row()

        # Extract values from model (adjust column names as needed)
        name = self.model2.data(self.model2.index(row_index, self.model2.fieldIndex("name_surname")))
        amount = self.model2.data(self.model2.index(row_index, self.model2.fieldIndex("amount")))
        date = self.model2.data(self.model2.index(row_index, self.model2.fieldIndex("payment_date")))

        # Format as HTML (you can style this)
        html = f"""
        <h2>Payment Receipt</h2>
        <p><b>Name:</b> {name}</p>
        <p><b>Amount:</b> {amount}</p>
        <p><b>Date:</b> {date}</p>
        """

        # Create document and print
        doc = QTextDocument()
        doc.setHtml(html)

        printer = QPrinter()

        dialog = QPrintDialog(printer, self)
        if dialog.exec_() == QPrintDialog.Accepted:
            doc.print_(printer)


    def refresh_table(self):
        self.model2.setFilter("")  # Clears filter
        self.model2.select()  # This reloads the data from DB

    def search_by_date(self):
        from_date_str = self.from_date.date().toString("yyyy-MM-dd HH:mm:ss")
        to_date_str = self.to_date.date().toString("yyyy-MM-dd HH:mm:ss")

        if self.contract_date_radio.isChecked():
            date_column = "date_of_C_O"
        elif self.closing_date_radio.isChecked():
            date_column = "date_of_addition"
        else:
            QMessageBox.warning(self, "შეცდომა", "აირჩიეთ თარიღის ტიპი.")
            return

        # Assuming your table has a date column named 'date'
        filter_str = f"{date_column} >= '{from_date_str}' AND {date_column} <= '{to_date_str}'"
        self.model2.setFilter(filter_str)
        self.model2.select()

    def apply_text_filter(self, text):
        column = ""

        if self.contract_radio.isChecked():
            column = "contract_id"
        elif self.name_radio.isChecked():
            column = "name_surname"
        elif self.id_radio.isChecked():
            column = "id_number"
        elif self.model_radio.isChecked():
            column = "model"
        elif self.tel_radio.isChecked():
            column = "tel_number"
        elif self.imei_radio.isChecked():
            column = "IMEI"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model2.setFilter(filter_str)
        else:
            self.model2.setFilter("")  # No filter if nothing selected

        self.model2.select()



    def export_to_excel(self):

        row_count = self.model2.rowCount()
        col_count = self.model2.columnCount()

        headers = [self.model2.headerData(col, Qt.Horizontal) for col in range(col_count)]
        data = [
            [self.model2.data(self.model2.index(row, col)) for col in range(col_count)]
            for row in range(row_count)
        ]

        df = pd.DataFrame(data, columns=headers)

        try:
            temp_path = os.path.join(tempfile.gettempdir(), "temp_export.xlsx")
            df.to_excel(temp_path, index=False)

            # Open Excel file
            os.startfile(temp_path)  # Safer and native on Windows
        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", str(e))

    def open_outflow_in_registry_edit_money_control_window(self):
        selected_indexes = self.table2.selectionModel().selectedRows()

        if selected_indexes:
            row_index = selected_indexes[0].row()
            model = self.table2.model()
            record_id = model.data(model.index(row_index, model.fieldIndex("unique_id")))

            self.open_outflow_in_registry_edit_money_control_window = EditRegistryOutflowWindow(record_id)
            self.open_outflow_in_registry_edit_money_control_window.show()

        else:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ შესაცვლელი ჩანაწერი")


    def delete_selected_row(self):
        selected_indexes = self.table2.selectionModel().selectedRows()
        if not selected_indexes:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ წასაშლელი ჩანაწერი")
            return

        row_index = selected_indexes[0].row()
        model = self.table2.model()
        unique_id_index = model.fieldIndex("unique_id")
        contract_id_index = model.fieldIndex("contract_id")
        amount_index = model.fieldIndex("additional_amount")
        status_index = model.fieldIndex("status")
        date_index = model.fieldIndex("date_of_addition")

        if unique_id_index == -1:
            QMessageBox.critical(self, "შეცდომა", "ვერ მოიძებნა უნიკალური ID სვეტი")
            return

        unique_id = model.data(model.index(row_index, unique_id_index))
        contract_id = model.data(model.index(row_index, contract_id_index))
        amount = model.data(model.index(row_index, amount_index))
        status = model.data(model.index(row_index, status_index))
        date_of_addition = model.data(model.index(row_index, date_index))

        reply = QMessageBox.question(
            self, "დადასტურება", "ნამდვილად გსურთ ამ ჩანაწერის წაშლა?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            # Remove the row from the model
            model.removeRow(row_index)

            # 2. Delete from given_and_additional_database
            conn_given = sqlite3.connect("Databases/given_and_additional_database.db")
            cur_given = conn_given.cursor()
            cur_given.execute("""
                    DELETE FROM given_and_additional_database
                    WHERE contract_id = ? AND amount = ? AND status = ? AND date_of_outflow = ?
                """, (contract_id, amount, status, date_of_addition))
            conn_given.commit()
            conn_given.close()

            # 3. Delete from outflow_order
            conn_outflow = sqlite3.connect("Databases/outflow_order.db")
            cur_outflow = conn_outflow.cursor()
            cur_outflow.execute("""
                    DELETE FROM outflow_order
                    WHERE contract_id = ? AND amount = ? AND status = ? AND date = ?
                """, (contract_id, amount, status, date_of_addition))
            conn_outflow.commit()
            conn_outflow.close()

            # 4. Delete amount from active_contracts database
            conn = sqlite3.connect("Databases/active_contracts.db")
            cursor = conn.cursor()

            # Fetch old value
            cursor.execute("""
                            SELECT additional_amounts, given_money, percent FROM active_contracts WHERE id = ?
                        """, (contract_id,))
            row = cursor.fetchone()

            if row and row[0]:
                old_amount = float(row[0])
                given_money = float(row[1])
                percent = float(row[2])
            else:
                old_amount = 0.0  # Or whatever default makes sense

            new_additional_amounts = old_amount - amount
            new_added_percent = (given_money + new_additional_amounts) * percent / 100

            cursor.execute("""
                                                                UPDATE active_contracts SET
                                                                    additional_amounts = ?,
                                                                    added_percents = ?
                                                                WHERE id = ?
                                                            """, (
                new_additional_amounts,
                new_added_percent,
                contract_id
            ))
            conn.commit()
            conn.close()



            if model.submitAll():
                QMessageBox.information(self, "წარმატება", "ჩანაწერი წაიშალა")
            else:
                QMessageBox.critical(self, "შეცდომა", "ჩანაწერის წაშლა ვერ მოხერხდა")


    # --------------------------------------------page2functions-----------------------------------------------------

    def refresh_table_2(self):
        self.model3.setFilter("")  # Clears filter
        self.model3.select()  # Reload data from principal_inflows.db

    def search_by_date_2(self):
        from_date_str = self.from_date_2.date().toString("yyyy-MM-dd HH:mm:ss")
        to_date_str = self.to_date_2.date().toString("yyyy-MM-dd HH:mm:ss")

        if self.date_radio_2.isChecked():
            date_column = "date_C_O"
        elif self.closing_date_radio_2.isChecked():
            date_column = "date_of_payment"
        else:
            QMessageBox.warning(self, "შეცდომა", "აირჩიეთ თარიღის ტიპი.")
            return

        filter_str = f"{date_column} >= '{from_date_str}' AND {date_column} <= '{to_date_str}'"
        self.model3.setFilter(filter_str)
        self.model3.select()

    def apply_text_filter_2(self, text):
        column = ""

        if self.name_radio_2.isChecked():
            column = "name_surname"
        elif self.id_radio_2.isChecked():
            column = "id_number"
        elif self.contract_radio_2.isChecked():
            column = "contract_id"
        elif self.model_radio_2.isChecked():
            column = "model"
        elif self.tel_radio_2.isChecked():
            column = "tel_number"
        elif self.imei_radio_2.isChecked():
            column = "IMEI"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model3.setFilter(filter_str)
        else:
            self.model3.setFilter("")

        self.model3.select()

    def export_to_excel_2(self):
        row_count = self.model3.rowCount()
        col_count = self.model3.columnCount()

        headers = [self.model3.headerData(col, Qt.Horizontal) for col in range(col_count)]
        data = [
            [self.model3.data(self.model3.index(row, col)) for col in range(col_count)]
            for row in range(row_count)
        ]

        df = pd.DataFrame(data, columns=headers)

        try:
            temp_path = os.path.join(tempfile.gettempdir(), "temp_export_page2.xlsx")
            df.to_excel(temp_path, index=False)
            os.startfile(temp_path)
        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", str(e))


    def open_edit_money_control_window_2(self):
        selected_indexes = self.table3.selectionModel().selectedRows()
        if selected_indexes:
            row_index = selected_indexes[0].row()
            model = self.table3.model()
            record_id = model.data(model.index(row_index, model.fieldIndex("unique_id")))
            self.open_edit_money_control_window_2 = EditInPrincipalInflowsInRegistryWindow(record_id)
            self.open_edit_money_control_window_2.show()
        else:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ შესაცვლელი ჩანაწერი")


    def delete_selected_row_2(self):
        selected_indexes = self.table3.selectionModel().selectedRows()
        if not selected_indexes:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ წასაშლელი ჩანაწერი")
            return

        row_index = selected_indexes[0].row()
        model_page2 = self.table3.model()
        unique_id_index = model_page2.fieldIndex("unique_id")
        contract_id_index = model_page2.fieldIndex("contract_id")
        principle_index = model_page2.fieldIndex("payment_amount")
        date_of_payment_index = model_page2.fieldIndex("date_of_payment")
        status_of_payment_index = model_page2.fieldIndex("status")


        if unique_id_index == -1:
            QMessageBox.critical(self, "შეცდომა", "ვერ მოიძებნა უნიკალური ID სვეტი")
            return

        unique_id = model_page2.data(model_page2.index(row_index, unique_id_index))
        contract_id_page2 = model_page2.data(model_page2.index(row_index, contract_id_index))
        principle_page2 = model_page2.data(model_page2.index(row_index, principle_index))
        date_of_payment_page2 = model_page2.data(model_page2.index(row_index, date_of_payment_index))
        status_of_payment_page2 = model_page2.data(model_page2.index(row_index, status_of_payment_index))




        reply = QMessageBox.question(
            self, "დადასტურება", "ნამდვილად გსურთ ამ ჩანაწერის წაშლა?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            model_page2.removeRow(row_index)

            conn = sqlite3.connect("Databases/active_contracts.db")
            cursor = conn.cursor()

            cursor.execute("SELECT principal_paid FROM active_contracts WHERE id = ?", (contract_id_page2,))
            row = cursor.fetchone()
            conn.close()

            if row:
                paid_principles_before = row[0]

            new_principles_amount = paid_principles_before - principle_page2

            conn = sqlite3.connect("Databases/active_contracts.db")
            cursor = conn.cursor()

            cursor.execute("""
                              UPDATE active_contracts SET
                                principal_paid = ?
                              WHERE id = ?
                          """, (
                new_principles_amount,
                contract_id_page2,
            ))
            conn.commit()
            conn.close()

            # Deleting from paid_principle_and_paid_percentage_database
            conn = sqlite3.connect("Databases/paid_principle_and_paid_percentage_database.db")
            cur_given = conn.cursor()
            cur_given.execute("""
                                            DELETE FROM paid_principle_and_paid_percentage_database
                                            WHERE contract_id = ? AND status = ? AND date_of_inflow = ?
                                        """, (contract_id_page2, status_of_payment_page2, date_of_payment_page2))
            conn.commit()
            conn.close()

            # Deleting from inflow_order_only_principal_amount
            conn = sqlite3.connect("Databases/inflow_order_only_principal_amount.db")
            cur_given = conn.cursor()
            cur_given.execute("""
                               DELETE FROM inflow_order_only_principal_amount
                               WHERE contract_id = ? AND payment_date = ?
                              """,
                              (contract_id_page2, date_of_payment_page2))
            conn.commit()
            conn.close()

            # Deleting from inflow_order_in_both
            conn = sqlite3.connect("Databases/inflow_order_both.db")
            cur_given = conn.cursor()
            cur_given.execute("""
                                           DELETE FROM inflow_order_both
                                           WHERE contract_id = ? AND payment_date = ? AND percent_paid_amount = ?
                                          """,
                              (contract_id_page2, date_of_payment_page2, 0))
            conn.commit()
            conn.close()




            if model_page2.submitAll():
                QMessageBox.information(self, "წარმატება", "ჩანაწერი წაიშალა")
            else:
                QMessageBox.critical(self, "შეცდომა", "ჩანაწერის წაშლა ვერ მოხერხდა")



    # --------------------------------------------page3functions-----------------------------------------------------
    def refresh_table_3(self):
        self.model3_1.setFilter("")
        self.model3_1.select()

    def search_by_date_3(self):
        from_date_str = self.from_date_3.date().toString("yyyy-MM-dd HH:mm:ss")
        to_date_str = self.to_date_3.date().toString("yyyy-MM-dd HH:mm:ss")

        if self.date_radio_3.isChecked():
            date_column = "date_of_C_O"
        elif self.closing_date_radio_3.isChecked():
            date_column = "date_of_percent_addition"
        else:
            QMessageBox.warning(self, "შეცდომა", "აირჩიეთ თარიღის ტიპი.")
            return

        filter_str = f"{date_column} >= '{from_date_str}' AND {date_column} <= '{to_date_str}'"
        self.model3_1.setFilter(filter_str)
        self.model3_1.select()

    def apply_text_filter_3(self, text):
        column = ""

        if self.name_radio_3.isChecked():
            column = "name_surname"
        elif self.id_radio_3.isChecked():
            column = "id_number"
        elif self.contract_radio_3.isChecked():
            column = "contract_id"
        elif self.model_radio_3.isChecked():
            column = "model"
        elif self.tel_radio_3.isChecked():
            column = "tel_number"
        elif self.imei_radio_3.isChecked():
            column = "IMEI"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model3_1.setFilter(filter_str)
        else:
            self.model3_1.setFilter("")

        self.model3_1.select()

    def export_to_excel_3(self):
        row_count = self.model3_1.rowCount()
        col_count = self.model3_1.columnCount()

        headers = [self.model3_1.headerData(col, Qt.Horizontal) for col in range(col_count)]
        data = [
            [self.model3_1.data(self.model3_1.index(row, col)) for col in range(col_count)]
            for row in range(row_count)
        ]

        df = pd.DataFrame(data, columns=headers)

        try:
            temp_path = os.path.join(tempfile.gettempdir(), "temp_export_page3.xlsx")
            df.to_excel(temp_path, index=False)
            os.startfile(temp_path)
        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", str(e))

    def open_edit_money_control_window_3(self):
        selected_indexes = self.table3_1.selectionModel().selectedRows()
        if selected_indexes:
            row_index = selected_indexes[0].row()
            model = self.table3_1.model()
            record_id = model.data(model.index(row_index, model.fieldIndex("unique_id")))
            self.edit_window_3 = EditAddingPercentWindow(record_id)
            self.edit_window_3.show()

        else:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ შესაცვლელი ჩანაწერი")

    def delete_selected_row_3(self):
        selected_indexes = self.table3_1.selectionModel().selectedRows()
        if not selected_indexes:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ წასაშლელი ჩანაწერი")
            return

        row_index = selected_indexes[0].row()
        model_3_1 = self.table3_1.model()
        unique_id_index = model_3_1.fieldIndex("unique_id")
        percent_amount_index = model_3_1.fieldIndex("percent_amount")
        contract_id_index = model_3_1.fieldIndex("contract_id")

        percent_amount_page3 = model_3_1.data(model_3_1.index(row_index, percent_amount_index))
        contract_id_page3 = model_3_1.data(model_3_1.index(row_index, contract_id_index))


        if unique_id_index == -1:
            QMessageBox.critical(self, "შეცდომა", "ვერ მოიძებნა უნიკალური ID სვეტი")
            return

        reply = QMessageBox.question(
            self, "დადასტურება", "ნამდვილად გსურთ ჩანაწერის წაშლა?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            model_3_1.removeRow(row_index)

            conn = sqlite3.connect("Databases/active_contracts.db")
            cursor = conn.cursor()

            cursor.execute("SELECT added_percents FROM active_contracts WHERE id = ?", (contract_id_page3,))
            row = cursor.fetchone()
            conn.close()

            if row:
                added_percents_before = row[0]
            else:
                print("No record found")


            new_percents_amount = added_percents_before - percent_amount_page3

            conn = sqlite3.connect("Databases/active_contracts.db")
            cursor = conn.cursor()

            cursor.execute("""
                                        UPDATE active_contracts SET
                                            added_percents = ?
                                        WHERE id = ?
                                    """, (
                new_percents_amount,
                contract_id_page3,
            ))
            conn.commit()
            conn.close()


            if model_3_1.submitAll():
                QMessageBox.information(self, "წარმატება", "ჩანაწერი წაიშალა")
            else:
                QMessageBox.critical(self, "შეცდომა", "წაშლა ვერ მოხერხდა")


    # --------------------------------------------page4functions-----------------------------------------------------

    def refresh_table_4(self):
        self.model4.setFilter("")
        self.model4.select()

    def search_by_date_4(self):
        from_date_str = self.from_date_4.date().toString("yyyy-MM-dd HH:mm:ss")
        to_date_str = self.to_date_4.date().toString("yyyy-MM-dd HH:mm:ss")

        if self.date_radio_4.isChecked():
            date_column = "date_of_C_O"
        elif self.closing_date_radio_4.isChecked():
            date_column = "date_of_percent_addition"
        else:
            QMessageBox.warning(self, "შეცდომა", "აირჩიეთ თარიღის ტიპი.")
            return

        filter_str = f"{date_column} >= '{from_date_str}' AND {date_column} <= '{to_date_str}'"
        self.model4.setFilter(filter_str)
        self.model4.select()

    def apply_text_filter_4(self, text):
        column = ""

        if self.name_radio_4.isChecked():
            column = "name_surname"
        elif self.id_radio_4.isChecked():
            column = "id_number"
        elif self.contract_radio_4.isChecked():
            column = "contract_id"
        elif self.model_radio_4.isChecked():
            column = "model"
        elif self.tel_radio_4.isChecked():
            column = "tel_number"
        elif self.imei_radio_4.isChecked():
            column = "IMEI"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model4.setFilter(filter_str)
        else:
            self.model4.setFilter("")

        self.model4.select()

    def export_to_excel_4(self):
        row_count = self.model4.rowCount()
        col_count = self.model4.columnCount()

        headers = [self.model4.headerData(col, Qt.Horizontal) for col in range(col_count)]
        data = [
            [self.model4.data(self.model4.index(row, col)) for col in range(col_count)]
            for row in range(row_count)
        ]

        df = pd.DataFrame(data, columns=headers)

        try:
            temp_path = os.path.join(tempfile.gettempdir(), "temp_export_page4.xlsx")
            df.to_excel(temp_path, index=False)
            os.startfile(temp_path)
        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", str(e))

    def open_edit_money_control_window_4(self):
        selected_indexes = self.table4.selectionModel().selectedRows()
        if selected_indexes:
            row_index = selected_indexes[0].row()
            model = self.table4.model()
            record_id = model.data(model.index(row_index, model.fieldIndex("unique_id")))
            self.edit_window_4 = EditPaidPercentWindow(record_id)
            self.edit_window_4.show()
        else:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ შესაცვლელი ჩანაწერი")

    def delete_selected_row_4(self):
        selected_indexes = self.table4.selectionModel().selectedRows()
        if not selected_indexes:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ წასაშლელი ჩანაწერი")
            return

        row_index = selected_indexes[0].row()
        model4 = self.table4.model()

        unique_id_index = model4.fieldIndex("unique_id")
        contract_id_index = model4.fieldIndex("contract_id")
        paid_percent_index = model4.fieldIndex("paid_amount")
        date_index = model4.fieldIndex("date_of_percent_addition")
        status_index = model4.fieldIndex("status")

        if unique_id_index == -1:
            QMessageBox.critical(self, "შეცდომა", "ვერ მოიძებნა უნიკალური ID სვეტი")
            return

        contract_id = model4.data(model4.index(row_index, contract_id_index))
        paid_percent_amount = model4.data(model4.index(row_index, paid_percent_index))
        percent_payment_date = model4.data(model4.index(row_index, date_index))
        status = model4.data(model4.index(row_index, status_index))

        reply = QMessageBox.question(
            self, "დადასტურება", "ნამდვილად გსურთ ჩანაწერის წაშლა?",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )

        if reply == QMessageBox.Yes:
            model4.removeRow(row_index)

            conn = sqlite3.connect("Databases/active_contracts.db")
            cursor = conn.cursor()

            cursor.execute("SELECT paid_percents FROM active_contracts WHERE id = ?", (contract_id,))
            row = cursor.fetchone()
            conn.close()

            if row:
                paid_percents_before = row[0]

            new_percents_amount = paid_percents_before - paid_percent_amount

            conn = sqlite3.connect("Databases/active_contracts.db")
            cursor = conn.cursor()

            cursor.execute("""
                                          UPDATE active_contracts SET
                                            paid_percents = ?
                                          WHERE id = ?
                                      """, (
                new_percents_amount,
                contract_id,
            ))
            conn.commit()
            conn.close()

            conn = sqlite3.connect("Databases/paid_principle_and_paid_percentage_database.db")
            cur_given = conn.cursor()
            cur_given.execute("""
                                   DELETE FROM paid_principle_and_paid_percentage_database
                                   WHERE contract_id = ? AND status = ? AND date_of_inflow = ?
                               """,
                              (contract_id, status, percent_payment_date))
            conn.commit()
            conn.close()

            conn = sqlite3.connect("Databases/inflow_order_only_percent_amount.db")
            cur_given = conn.cursor()
            cur_given.execute("""
                                               DELETE FROM inflow_order_only_percent_amount
                                               WHERE contract_id = ? AND payment_date = ?
                                           """,
                              (contract_id, percent_payment_date))
            conn.commit()
            conn.close()

            # Deleting from inflow_order_in_both
            conn = sqlite3.connect("Databases/inflow_order_both.db")
            cur_given = conn.cursor()
            cur_given.execute("""
                                  DELETE FROM inflow_order_both
                                  WHERE contract_id = ? AND payment_date = ? AND principle_paid_amount = ?
                              """,
                              (contract_id, percent_payment_date, 0))
            conn.commit()
            conn.close()


            if model4.submitAll():
                QMessageBox.information(self, "წარმატება", "ჩანაწერი წაიშალა")
            else:
                QMessageBox.critical(self, "შეცდომა", "წაშლა ვერ მოხერხდა")

    # --------------------------------------------page5functions-----------------------------------------------------
    def show_table5_context_menu(self, position: QPoint):
        index = self.table5.indexAt(position)
        if not index.isValid():
            return

        menu = QMenu()
        print_action = QAction(" ამობეჭდვა ", self)
        print_action.setIcon(QIcon("Icons/printer_icon.png"))
        print_action.triggered.connect(self.print_table5_selected_row)
        menu.addAction(print_action)
        menu.exec_(self.table5.viewport().mapToGlobal(position))


    def print_table5_selected_row(self):
        selected = self.table5.selectionModel().selectedRows()
        if not selected:
            print("No row selected.")
            return

        row_index = selected[0].row()
        contract_id = self.model5.data(self.model5.index(row_index, self.model5.fieldIndex("contract_id")))
        name = self.model5.data(self.model5.index(row_index, self.model5.fieldIndex("name_surname")))
        date_raw = self.model5.data(self.model5.index(row_index, self.model5.fieldIndex("date")))
        unique_id = self.model5.data(self.model5.index(row_index, self.model5.fieldIndex("unique_id")))
        additional_amount = self.model5.data(self.model5.index(row_index, self.model5.fieldIndex("amount")))
        dt = datetime.strptime(date_raw, "%Y-%m-%d %H:%M:%S")
        date = dt.strftime("%d-%m-%Y")

        replacements = {
            '{name_surname}': name or "",
            '{additional_amount}': str(additional_amount) if additional_amount is not None else "",
            '{date}': date or "",
            '{contract_id}': str(contract_id or ""),
            '{unique_id}': unique_id or "",
            '{organization_name}': getattr(self, "organisation", ""),
        }

        def replace_in_paragraph(paragraph, replacements):
            full_text = ''.join(run.text for run in paragraph.runs)
            new_text = full_text
            for key, value in replacements.items():
                new_text = new_text.replace(key, str(value))

            if new_text != full_text:
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

        # Load the Word template
        doc = Document("Templates/additional_money_template.docx")

        # Replace in normal paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)

        # Replace in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)

        # 3. Save new doc
        # Ensure folder exists
        output_dir = "GeneratedContracts"
        os.makedirs(output_dir, exist_ok=True)

        # Construct file name
        output_filename = f"outflow_order_{unique_id}_{contract_id}_{name}.docx"
        output_path = os.path.join(output_dir, output_filename)

        # Save document
        doc.save(output_path)

        # Open in Word and wait
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = True
            word_doc = word.Documents.Open(os.path.abspath(output_path))


        except Exception as e:
            print("Error:", e)
            print("Document saved at:", output_path)

        # # 4. Optional: Print using MS Word (Windows only)
        # try:
        #     word = win32com.client.Dispatch("Word.Application")
        #     word.Visible = False
        #     word.Documents.Open(os.path.abspath(output_path)).PrintOut()
        #     word.Quit()
        # except Exception as e:
        #     print("Printing failed:", e)
        #     # fallback: open Word file manually
        #     os.startfile(output_path)



    def apply_text_filter_5(self, text):
        column = ""

        if self.name_radio_5.isChecked():
            column = "name_surname"
        elif self.contract_radio_5.isChecked():
            column = "contract_id"
        elif self.tel_radio_5.isChecked():
            column = "tel_number"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model5.setFilter(filter_str)
        else:
            self.model5.setFilter("")

        self.model5.select()

    # --------------------------------------------page6functions-----------------------------------------------------
    def show_table6_1_context_menu(self, position):
        index = self.table6_1.indexAt(position)
        if not index.isValid():
            return

        menu = QMenu()
        print_action = QAction(" ამობეჭდვა ", self)
        print_action.setIcon(QIcon("Icons/printer_icon.png"))
        print_action.triggered.connect(self.print_table6_1_selected_row)
        menu.addAction(print_action)
        menu.exec_(self.table6_1.viewport().mapToGlobal(position))

    def print_table6_1_selected_row(self):
        selected = self.table6_1.selectionModel().selectedRows()
        if not selected:
            print("No row selected.")
            return

        row_index = selected[0].row()
        name = self.model6_1.data(self.model6_1.index(row_index, self.model6_1.fieldIndex("name_surname")))
        amount = self.model6_1.data(self.model6_1.index(row_index, self.model6_1.fieldIndex("amount")))
        date = self.model6_1.data(self.model6_1.index(row_index, self.model6_1.fieldIndex("payment_date")))

        html = f"""
        <h2>გადახდის ქვითარი</h2>
        <p><b>სახელი:</b> {name}</p>
        <p><b>თანხა:</b> {amount}</p>
        <p><b>თარიღი:</b> {date}</p>
        """
        doc = QTextDocument()
        doc.setHtml(html)
        printer = QPrinter()
        dialog = QPrintDialog(printer, self)
        if dialog.exec_() == QPrintDialog.Accepted:
            doc.print_(printer)

    def show_table6_2_context_menu(self, position):
        index = self.table6_2.indexAt(position)
        if not index.isValid():
            return

        menu = QMenu()
        print_action = QAction(" ამობეჭდვა ", self)
        print_action.setIcon(QIcon("Icons/printer_icon.png"))
        print_action.triggered.connect(self.print_table6_2_selected_row)
        menu.addAction(print_action)
        menu.exec_(self.table6_2.viewport().mapToGlobal(position))

    def print_table6_2_selected_row(self):
        selected = self.table6_2.selectionModel().selectedRows()
        if not selected:
            print("No row selected.")
            return

        row_index = selected[0].row()
        name = self.model6_2.data(self.model6_2.index(row_index, self.model6_2.fieldIndex("name_surname")))
        amount = self.model6_2.data(self.model6_2.index(row_index, self.model6_2.fieldIndex("amount")))
        date = self.model6_2.data(self.model6_2.index(row_index, self.model6_2.fieldIndex("payment_date")))

        html = f"""
        <h2>გადახდის ქვითარი</h2>
        <p><b>სახელი:</b> {name}</p>
        <p><b>თანხა:</b> {amount}</p>
        <p><b>თარიღი:</b> {date}</p>
        """
        doc = QTextDocument()
        doc.setHtml(html)
        printer = QPrinter()
        dialog = QPrintDialog(printer, self)
        if dialog.exec_() == QPrintDialog.Accepted:
            doc.print_(printer)

    def show_table6_3_context_menu(self, position):
        index = self.table6_3.indexAt(position)
        if not index.isValid():
            return

        menu = QMenu()
        print_action = QAction(" ამობეჭდვა ", self)
        print_action.setIcon(QIcon("Icons/printer_icon.png"))
        print_action.triggered.connect(self.print_table6_3_selected_row)
        menu.addAction(print_action)
        menu.exec_(self.table6_3.viewport().mapToGlobal(position))

    def print_table6_3_selected_row(self):
        selected = self.table6_3.selectionModel().selectedRows()
        if not selected:
            print("No row selected.")
            return

        row_index = selected[0].row()
        name = self.model6_3.data(self.model6_3.index(row_index, self.model6_3.fieldIndex("name_surname")))
        amount = self.model6_3.data(self.model6_3.index(row_index, self.model6_3.fieldIndex("amount")))
        date = self.model6_3.data(self.model6_3.index(row_index, self.model6_3.fieldIndex("payment_date")))

        html = f"""
        <h2>გადახდის ქვითარი</h2>
        <p><b>სახელი:</b> {name}</p>
        <p><b>თანხა:</b> {amount}</p>
        <p><b>თარიღი:</b> {date}</p>
        """
        doc = QTextDocument()
        doc.setHtml(html)
        printer = QPrinter()
        dialog = QPrintDialog(printer, self)
        if dialog.exec_() == QPrintDialog.Accepted:
            doc.print_(printer)

    def apply_text_filter_6_1(self, text):
        column = ""

        if self.contract_radio_6_1.isChecked():
            column = "contract_id"
        elif self.name_radio_6_1.isChecked():
            column = "name_surname"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model6_1.setFilter(filter_str)
        else:
            self.model6_1.setFilter("")

        self.model6_1.select()


    def apply_text_filter_6_2(self, text):
        column = ""

        if self.contract_radio_6_2.isChecked():
            column = "contract_id"
        elif self.name_radio_6_2.isChecked():
            column = "name_surname"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model6_2.setFilter(filter_str)
        else:
            self.model6_2.setFilter("")

        self.model6_2.select()


    def apply_text_filter_6_3(self, text):
        column = ""

        if self.contract_radio_6_3.isChecked():
            column = "contract_id"
        elif self.name_radio_6_3.isChecked():
            column = "name_surname"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model6_3.setFilter(filter_str)
        else:
            self.model6_3.setFilter("")

        self.model6_3.select()

    # --------------------------------------------page7functions-----------------------------------------------------
    def show_table7_context_menu(self, position):
        index = self.table7.indexAt(position)
        if not index.isValid():
            return

        menu = QMenu()
        print_action = QAction(" ამობეჭდვა ", self)
        print_action.setIcon(QIcon("Icons/printer_icon.png"))
        print_action.triggered.connect(self.print_table7_selected_row)
        menu.addAction(print_action)
        menu.exec_(self.table7.viewport().mapToGlobal(position))

    def print_table7_selected_row(self):
        selected = self.table7.selectionModel().selectedRows()
        if not selected:

            return

        row_index = selected[0].row()
        contract_id = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("contract_id")))
        name = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("name_surname")))
        given_money = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("given_money")))
        date_raw = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("contract_open_date")))
        id_number = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("id_number")))
        item_name = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("item_name")))
        model = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("model")))
        imei = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("IMEI")))
        comment = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("comment")))
        trusted_person = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("trusted_person")))
        tel_number = self.model7.data(self.model7.index(row_index, self.model7.fieldIndex("tel_number")))
        dt = datetime.strptime(date_raw, "%Y-%m-%d %H:%M:%S")
        date = dt.strftime("%d-%m-%Y")

        replacements = {
            '{name_surname}': name or "",
            '{given_money}': str(given_money) if given_money is not None else "",
            '{date}': date or "",
            '{contract_id}': str(contract_id or ""),
            '{id_number}': id_number or "",
            '{IMEI}': imei or "",
            '{model}': model or "",
            '{item_name}': item_name or "",
            '{comment}': comment or "",
            '{trusted_person}': trusted_person or "",
            '{tel_number}': tel_number or "",
            '{organization_name}': getattr(self, "organisation", ""),
            '{operator_name}': getattr(self, "name_of_user", "")
        }

        def replace_in_paragraph(paragraph, replacements):
            full_text = ''.join(run.text for run in paragraph.runs)
            new_text = full_text
            for key, value in replacements.items():
                new_text = new_text.replace(key, str(value))

            if new_text != full_text:
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

        # Load the Word template
        doc = Document("Templates/contract_template.docx")

        # Replace in normal paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)

        # Replace in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)

        # 3. Save new doc
        # Ensure folder exists
        output_dir = "GeneratedContracts"
        os.makedirs(output_dir, exist_ok=True)

        # Construct file name
        output_filename = f"contract_{contract_id}_{name}.docx"
        output_path = os.path.join(output_dir, output_filename)

        # Save document
        doc.save(output_path)

        # Open in Word and wait
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = True
            word_doc = word.Documents.Open(os.path.abspath(output_path))


        except Exception as e:
            print("Error:", e)
            print("Document saved at:", output_path)

        # # 4. Optional: Print using MS Word (Windows only)
        # try:
        #     word = win32com.client.Dispatch("Word.Application")
        #     word.Visible = False
        #     word.Documents.Open(os.path.abspath(output_path)).PrintOut()
        #     word.Quit()
        # except Exception as e:
        #     print("Printing failed:", e)
        #     # fallback: open Word file manually
        #     os.startfile(output_path)

    def apply_text_filter_7(self, text):
        column = ""

        if self.contract_radio_7.isChecked():
            column = "contract_id"
        elif self.name_radio_7.isChecked():
            column = "name_surname"
        elif self.id_radio_7.isChecked():
            column = "id_number"
        elif self.model_radio_7.isChecked():
            column = "model"
        elif self.tel_radio_7.isChecked():
            column = "tel_number"
        elif self.imei_radio_7.isChecked():
            column = "IMEI"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model7.setFilter(filter_str)
        else:
            self.model7.setFilter("")

        self.model7.select()