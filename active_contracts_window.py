import sqlite3
import sys
import time
from datetime import datetime, timedelta
from email.policy import default

import win32com.client

from PyQt5.QtCore import Qt, QSize, QDate, QDateTime
from PyQt5.QtGui import QIcon, QBrush, QColor, QTextDocument
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
from PyQt5.QtSql import QSqlDatabase, QSqlTableModel
from PyQt5.QtWidgets import QWidget, QGridLayout, QGroupBox, QToolButton, QCheckBox, QLabel, QRadioButton, QButtonGroup, \
    QLineEdit, QDateEdit, QPushButton, QTableView, QAbstractItemView, QMessageBox, QMenu, QAction, QTableWidgetItem

from docx import Document
import tempfile
import pandas as pd
import os
import subprocess
from PyQt5.QtWidgets import QMessageBox

from add_money_window import AddMoney
from contract_color_delegate import ContractColorDelegate
from detail_window import DetailWindow
from open_edit_window import EditWindow
from open_add_window import AddWindow
from payment_confirm_window import PaymentConfirmWindow
from payment_window import PaymentWindow

class ActiveContracts(QWidget):
    def __init__(self, role, name_of_user, organisation, id_number_of_user):
        super().__init__()
        self.setWindowTitle("მოქმედი ხელშეკრულებები")
        self.initialize_active_contracts_database()
        self.initialize_contracts_database()
        self.initialize_closed_contracts_database()
        self.initialize_given_and_additional_database()
        self.initialize_paid_principle_and_paid_percentage_database()
        self.initialize_paid_principle_registry_database()
        self.initialize_outflow_order_database()
        self.initialize_outflow_in_registry_database()
        self.initialize_adding_percent_amount_database()
        self.initialize_paid_percent_amount_database()
        self.initialize_inflow_order_only_principal_amount_database()
        self.initialize_inflow_order_only_percent_amount_database()
        self.initialize_blk_list_database()
        self.initialize_inflow_order_both_database()



        self.setWindowIcon(QIcon("Icons/contract_icon.png"))
        self.resize(1400, 800)
        self.role = role
        self.name_of_user = name_of_user
        self.organisation = organisation
        self.id_number_of_user = id_number_of_user

        layout = QGridLayout()

        # --------------------------------------------Box1-----------------------------------------------------
        box1 = QGroupBox("ნავიგაცია")
        box1.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout = QGridLayout()

        # Box 1 buttons
        # Add_item
        add_item = QToolButton()
        add_item.setText(" დამატება ")
        add_item.setIcon(QIcon("Icons/add_data.png"))
        add_item.setIconSize(QSize(37, 40))
        add_item.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        add_item.setStyleSheet("font-size: 16px;")
        add_item.clicked.connect(self.open_add_window)

        # Edit item
        edit_item = QToolButton()
        edit_item.setText(" რედაქტირება ")
        edit_item.setIcon(QIcon("Icons/edit_data.png"))
        edit_item.setIconSize(QSize(37, 40))
        edit_item.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        edit_item.setStyleSheet("font-size: 16px;")
        edit_item.clicked.connect(self.open_edit_window)

        # Choosing fields
        choosing_the_field = QToolButton()
        choosing_the_field.setText(" ველების არჩევა ")
        choosing_the_field.setIcon(QIcon("Icons/data_modify.png"))
        choosing_the_field.setIconSize(QSize(37, 40))
        choosing_the_field.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        choosing_the_field.setStyleSheet("font-size: 16px;")
        # choosing_the_field.clicked.connect(self.open_help_window)

        # Export
        export = QToolButton()
        export.setText(" ექსპორტი ")
        export.setIcon(QIcon("Icons/excel_icon.png"))
        export.setIconSize(QSize(37, 40))
        export.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        export.setStyleSheet("font-size: 16px;")
        export.clicked.connect(self.export_to_excel)

        # Checkbox
        # all_together = QCheckBox("ყველა ერთად")

        # BlankUnderTheCheckBox
        # blank_under_the_check_box = QLabel("")
        # blank_under_the_check_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        # blank_under_the_check_box.setFixedHeight(30)


        # Box1 widgets
        box_layout.addWidget(add_item, 0, 0)
        box_layout.addWidget(edit_item, 0, 2)
        # box_layout.addWidget(choosing_the_field, 0, 2)
        box_layout.addWidget(export, 0, 3)
        # box_layout.addWidget(all_together, 1, 1)
        # box_layout.addWidget(blank_under_the_check_box, 1, 2)

        # Placing box1
        layout.addWidget(box1, 0, 0, 1, 1)
        box1.setLayout(box_layout)

        # --------------------------------------------Box2-----------------------------------------------------
        box2 = QGroupBox("ძებნა")
        box2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout2 = QGridLayout()

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("მოძებნე აქ")
        self.search_input.textChanged.connect(self.apply_text_filter)

        # Radio Button
        self.contract_radio = QRadioButton("ხელშეკრულების N")
        self.name_radio = QRadioButton("სახელი და გვარი")
        self.id_radio = QRadioButton("პირადი N")
        self.model_radio = QRadioButton("მოდელი")
        self.tel_radio = QRadioButton("ტელეფონის ნომერი")
        self.imei_radio = QRadioButton("IMEI")

        # Grouping buttons
        button_group = QButtonGroup()
        button_group.addButton(self.contract_radio)
        button_group.addButton(self.name_radio)
        button_group.addButton(self.id_radio)
        button_group.addButton(self.model_radio)
        button_group.addButton(self.tel_radio)
        button_group.addButton(self.imei_radio)

        # Making Widgets
        box_layout2.addWidget(self.contract_radio)
        box_layout2.addWidget(self.name_radio)
        box_layout2.addWidget(self.id_radio)
        box_layout2.addWidget(self.model_radio)
        box_layout2.addWidget(self.tel_radio)
        box_layout2.addWidget(self.imei_radio)
        box_layout2.addWidget(self.search_input, 2, 1)



        layout.addWidget(box2, 0, 1, 1, 1)
        box2.setLayout(box_layout2)

        # --------------------------------------------Box3-----------------------------------------------------
        box3 = QGroupBox("თარიღის მიხედვით გაფილტვრა")
        box3.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout3 = QGridLayout()

        self.from_date = QDateEdit()
        self.from_date.setCalendarPopup(True)
        self.from_date.setDate(QDate.currentDate())  # Today

        self.to_date = QDateEdit()
        self.to_date.setCalendarPopup(True)
        self.to_date.setDate(QDate.currentDate().addDays(1))  # Tomorrow
        #
        box_layout3.addWidget(QLabel("*დან თარიღი:"))
        box_layout3.addWidget(self.from_date)
        box_layout3.addWidget(QLabel("*მდე თარიღი:"))
        box_layout3.addWidget(self.to_date)

        filter_button = QPushButton("ძებნა")
        filter_button.clicked.connect(self.search_by_date)
        refresh_button = QPushButton("განახლება")
        refresh_button.clicked.connect(self.refresh_table)

        box_layout3.addWidget(refresh_button)
        box_layout3.addWidget(filter_button)

        layout.addWidget(box3, 0, 2, 1, 1)
        box3.setLayout(box_layout3)


        # --------------------------------------------Table-----------------------------------------------------
        self.db = QSqlDatabase.addDatabase("QSQLITE")
        self.db.setDatabaseName("Databases/active_contracts.db")
        if not self.db.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table = QTableView()
        self.model = QSqlTableModel(self, self.db)
        self.model.setTable("active_contracts_view")
        self.model.setFilter("is_visible = 'აქტიური'")
        self.model.select()

        # Block to rename columns
        record = self.model.record()
        column_indices = {record.field(i).name(): i for i in range(record.count())}

        column_labels = {
            "id": "ხელშეკრულების N",
            "date": "გაფორმების თარიღი",
            "days_after_C_O":"დღის რაოდენობა",
            "name_surname": "სახელი და გვარი",
            "id_number": "პირადი ნომერი",
            "tel_number": "ტელეფონის ნომერი",
            "item_name": "ნივთის დასახელება",
            "model": "მოდელი",
            "imei": "IMEI",
            "type": "დატოვების ტიპი",
            "trusted_person": "მინდობილი პირი",
            "comment": "კომენტარი",
            "given_money": "გაცემული ძირი თანხა",
            "percent": "პროცენტი",
            "day_quantity": "პროცენტის დღეების რაოდენობა",
            "additional_amounts": "დამატებული თანხები",
            "principal_paid": "გადახდილი ძირი თანხა",
            "principal_should_be_paid": "გადასახდელი ძირი თანხა",
            "added_percents": "დარიცხული პროცენტები",
            "paid_percents": "გადახდილი პროცენტები",
            "percent_should_be_paid": "გადასახდელი პროცენტები",
            "is_visible": "ხელშეკრულების სტატუსი"
        }

        for name, label in column_labels.items():
            if name in column_indices:
                self.model.setHeaderData(column_indices[name], Qt.Horizontal, label)

        # Continue as usual
        self.table.setModel(self.model)
        # Make header text bold
        header = self.table.horizontalHeader()
        font = header.font()
        font.setBold(True)
        header.setFont(font)
        header.setStyleSheet("""
            QHeaderView::section {
                padding: 4px 8px;
            }
        """)
        # Continue as usual
        delegate = ContractColorDelegate(self.table)
        self.table.setItemDelegate(delegate)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)  # Read-only table
        self.table.setSelectionBehavior(QTableView.SelectRows)
        self.table.setSelectionMode(QTableView.SingleSelection)
        self.table.clicked.connect(self.row_selected)
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_context_menu)
        self.table.doubleClicked.connect(self.open_detail_window)

        self.table.resizeColumnsToContents()
        layout.addWidget(self.table, 1, 0, 4, 5)



        # Create QLabel to display the total sums
        self.total_label = QLabel("")
        self.total_label.setStyleSheet("""
            background-color: #f7f3e9;
            font-weight: bold;
            font-size: 14px;
            padding: 2px;
            border: 1px solid gray;
        """)
        self.total_label.setFixedHeight(28)
        self.total_label.setAlignment(Qt.AlignCenter)

        # Add the label below the table
        layout.addWidget(self.total_label, 6, 1, 1, 3)

        self.update_summary_footer()

    # --------------------------------------------BelowBox1-----------------------------------------------------
        below_box1 = QGroupBox("სტატუსი და მინდობილობა")
        below_box1.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        below_box_layout1 = QGridLayout()


        self.blank1_below_box = QLabel("")
        self.blank1_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank1_below_box.setFixedHeight(30)

        self.blank2_below_box = QLabel("")
        self.blank2_below_box.setWordWrap(True)
        self.blank2_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank2_below_box.setFixedHeight(30)

        # Box1 widgets
        below_box_layout1.addWidget(self.blank1_below_box, 0, 0)
        below_box_layout1.addWidget(self.blank2_below_box, 1, 0)

        # Placing box1
        layout.addWidget(below_box1, 7, 0, 2, 1)
        below_box1.setLayout(below_box_layout1)

        # --------------------------------------------BelowBox2-----------------------------------------------------

        below_box2 = QGroupBox("კომენტარი და შემდეგი პროცენტი")
        below_box2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        below_box_layout2 = QGridLayout()

        self.blank3_below_box = QLabel("")
        self.blank3_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank3_below_box.setFixedHeight(30)

        self.blank4_below_box = QLabel("")
        self.blank4_below_box.setWordWrap(True)
        self.blank4_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank4_below_box.setFixedHeight(30)

        # BelowBox2 widgets
        below_box_layout2.addWidget(self.blank3_below_box, 0, 0)
        below_box_layout2.addWidget(self.blank4_below_box, 1, 0)

        # Placing BelowBox2
        layout.addWidget(below_box2, 7, 1, 2, 1)
        below_box2.setLayout(below_box_layout2)

        # --------------------------------------------BelowBox3-----------------------------------------------------
        below_box3 = QGroupBox("თანხის კონტროლი")
        below_box3.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        below_box_layout3 = QGridLayout()

        blank15_below_box = QLabel("გაცემული ძირი თანხა:")
        blank17_below_box = QLabel("დამატებითი თანხები:")
        blank18_below_box = QLabel("გადასახდელი %-ები:")
        blank19_below_box = QLabel("გადახდილი ძირი თანხა:")
        blank110_below_box = QLabel("გადახდილი %-ები:")
        blank111_below_box = QLabel("სულ გადასახდელია:")

        self.blank5_below_box = QLabel("")
        self.blank5_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank5_below_box.setFixedHeight(30)
        self.blank5_below_box.setFixedWidth(140)

        self.blank7_below_box = QLabel("")
        self.blank7_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank7_below_box.setFixedHeight(30)
        self.blank7_below_box.setFixedWidth(140)

        self.blank8_below_box = QLabel("")
        self.blank8_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank8_below_box.setFixedHeight(30)
        self.blank8_below_box.setFixedWidth(140)

        self.blank9_below_box = QLabel("")
        self.blank9_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank9_below_box.setFixedHeight(30)
        self.blank9_below_box.setFixedWidth(140)

        self.blank11_below_box = QLabel("")
        self.blank11_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank11_below_box.setFixedHeight(30)
        self.blank11_below_box.setFixedWidth(140)

        self.blank10_below_box = QLabel("")
        self.blank10_below_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        self.blank10_below_box.setFixedHeight(30)
        self.blank10_below_box.setFixedWidth(140)

        # BelowBox3 widgets
        below_box_layout3.addWidget(blank15_below_box, 0, 0)
        below_box_layout3.addWidget(blank17_below_box, 1, 0)
        below_box_layout3.addWidget(blank18_below_box, 2, 0)
        below_box_layout3.addWidget(self.blank5_below_box, 0, 1)
        below_box_layout3.addWidget(self.blank7_below_box, 1, 1)
        below_box_layout3.addWidget(self.blank8_below_box, 2, 1)
        below_box_layout3.addWidget(blank19_below_box, 0, 2)
        below_box_layout3.addWidget(blank110_below_box, 1, 2)
        below_box_layout3.addWidget(blank111_below_box, 2, 2)
        below_box_layout3.addWidget(self.blank9_below_box, 0, 3)
        below_box_layout3.addWidget(self.blank10_below_box, 1, 3)
        below_box_layout3.addWidget(self.blank11_below_box, 2, 3)

        # Placing BelowBox3
        layout.addWidget(below_box3, 7, 2, 2, 1)
        below_box3.setLayout(below_box_layout3)





        # --------------------------------------------Layout-----------------------------------------------------
        self.setLayout(layout)

        self.load_data()




        # --------------------------------------------Functions-----------------------------------------------------
    def open_add_window(self):
        self.open_add_window = AddWindow(self.organisation, self.name_of_user)
        self.open_add_window.show()

    def open_edit_window(self):
        selected_indexes = self.table.selectionModel().selectedRows()
        if selected_indexes:
            row_index = selected_indexes[0].row()
            model = self.table.model()
            record_id = model.data(model.index(row_index, model.fieldIndex("id")))
            self.open_edit_window = EditWindow(record_id, self.role)
            self.open_edit_window.show()
        else:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ შესაცვლელი ჩანაწერი")

    def refresh_table(self):
        self.model.setFilter("is_visible = 'აქტიური'")  # Clears filter
        self.model.select()  # This reloads the data from DB
        self.update_summary_footer()

    def search_by_date(self):
        from_date_str = self.from_date.date().toString("yyyy-MM-dd")
        to_date_str = self.to_date.date().toString("yyyy-MM-dd")

        # Assuming your table has a date column named 'date'
        filter_str = f"date >= '{from_date_str}' AND date <= '{to_date_str}'"
        self.model.setFilter(f"{filter_str} AND is_visible = 'აქტიური'")
        self.model.select()
        self.update_summary_footer()

    def apply_text_filter(self, text):
        column = ""

        if self.contract_radio.isChecked():
            column = "id"
        elif self.name_radio.isChecked():
            column = "name_surname"
        elif self.id_radio.isChecked():
            column = "id_number"
        elif self.model_radio.isChecked():
            column = "model"
        elif self.tel_radio.isChecked():
            column = "tel_number"
        elif self.imei_radio.isChecked():
            column = "imei"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model.setFilter(f"{filter_str} AND is_visible = 'აქტიური'")
        else:
            self.model.setFilter("is_visible = 'აქტიური'")  # No filter if nothing selected

        self.model.select()
        self.update_summary_footer()


    def row_selected(self, index):
        row = index.row()
        model = self.model  # QSqlTableModel

        def safe_data(field_name, default=None):
            col = model.fieldIndex(field_name)
            if col == -1:
                print(f"Warning: Column '{field_name}' not found")
                return default
            val = model.data(model.index(row, col))
            return val if val is not None else default

        # Read text/string fields safely
        status = safe_data("type", "")
        delegate_person = safe_data("trusted_person", "")
        comment = safe_data("comment", "")

        # Read date string and parse
        date_str = safe_data("date", "")
        contract_datetime = QDateTime.fromString(date_str, "yyyy-MM-dd HH:mm:ss")
        if not contract_datetime.isValid():
            print(f"Invalid date string: '{date_str}' — using current date as fallback")
            contract_datetime = QDateTime.currentDateTime()

        # Read day_quantity safely and convert to int
        try:
            day_quantity = int(safe_data("day_quantity"))
            days_after_c_o = int(safe_data("days_after_C_O"))
        except (TypeError, ValueError):
            day_quantity = 0
            days_after_c_o = 0

        # Calculate next payment date based on day_quantity
        next_payment_date = contract_datetime.addDays(day_quantity - 1) if day_quantity >= 0 else contract_datetime

        # Move next_payment_date forward if already passed
        today = QDate.currentDate()
        days_on_excess = days_after_c_o % day_quantity
        days_to_add = day_quantity - days_on_excess
        while next_payment_date.date() < today and day_quantity > 0:
                next_payment_date = next_payment_date.addDays(days_to_add)

        next_payment_str = next_payment_date.toString("dd.MM.yyyy")

        # Numeric fields with safe float conversion
        def safe_float(field_name):
            try:
                return float(safe_data(field_name, 0))
            except (TypeError, ValueError):
                return 0.0

        principal_given = safe_float("given_money")
        additional_amounts = safe_float("additional_amounts")
        percent_should_be_paid = safe_float("percent_should_be_paid")
        principal_paid = safe_float("principal_paid")
        percent_paid = safe_float("paid_percents")
        principal_should_be_paid = safe_float("principal_should_be_paid")
        total_due = percent_should_be_paid + principal_should_be_paid

        # Update your UI labels safely
        self.blank1_below_box.setText(str(status))
        self.blank2_below_box.setText(str(delegate_person))
        self.blank3_below_box.setText(str(comment))
        self.blank4_below_box.setText(next_payment_str)
        self.blank5_below_box.setText(f"{principal_given:.2f}")
        self.blank7_below_box.setText(f"{additional_amounts:.2f}")
        self.blank8_below_box.setText(f"{percent_should_be_paid:.2f}")
        self.blank9_below_box.setText(f"{principal_paid:.2f}")
        self.blank10_below_box.setText(f"{percent_paid:.2f}")
        self.blank11_below_box.setText(f"{total_due:.2f}")

    def export_to_excel(self):

        row_count = self.model.rowCount()
        col_count = self.model.columnCount()

        headers = [self.model.headerData(col, Qt.Horizontal) for col in range(col_count)]
        data = [
            [self.model.data(self.model.index(row, col)) for col in range(col_count)]
            for row in range(row_count)
        ]

        df = pd.DataFrame(data, columns=headers)

        try:
            temp_path = os.path.join(tempfile.gettempdir(), "temp_export.xlsx")
            df.to_excel(temp_path, index=False)

            # Open Excel file
            os.startfile(temp_path)  # Safer and native on Windows
        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", str(e))

    # Creating right click menu
    def show_context_menu(self, position):
        indexes = self.table.selectedIndexes()
        if not indexes:
            return  # No row is selected

        menu = QMenu()

        print_action = QAction(" ხელშეკრულების ამობეჭდვა ", self)
        print_action.setIcon(QIcon("Icons/printer_icon.png"))
        print_action.triggered.connect(self.print_selected_row_page)


        add_money_action = QAction(" თანხის დამატება ", self)
        add_money_action.setIcon(QIcon("Icons/add_money.png"))
        add_money_action.triggered.connect(self.add_money_selected_row)

        pay_percent_action = QAction(" პროცენტის ან ძირი თანხის გადახდა ", self)
        pay_percent_action.setIcon(QIcon("Icons/percent_payment_icon.png"))
        pay_percent_action.triggered.connect(self.pay_percent_selected_row)

        closing_contract_action = QAction(" ხელშეკრულების დახურვა ", self)
        closing_contract_action.setIcon(QIcon("Icons/closed_contracts.png"))
        closing_contract_action.triggered.connect(self.closing_contract_selected_row)

        add_to_the_black_list_action = QAction(" შავ სიაში შეტანა ", self)
        add_to_the_black_list_action.setIcon(QIcon("Icons/blacklist.png"))
        add_to_the_black_list_action.triggered.connect(self.add_to_the_black_list_selected_row)

        # transfer_to_selling_items_action = QAction("", self)
        # transfer_to_selling_items_action.setIcon(QIcon("Icons/sell_icon.png"))
        # transfer_to_selling_items_action.triggered.connect(self.transfer_to_selling_items_selected_row)

        # got_back_from_police_action = QAction("", self)
        # got_back_from_police_action.setIcon(QIcon("Icons/police-badge.png"))
        # got_back_from_police_action.triggered.connect(self.got_back_from_police_action_selected_row)


        menu.addAction(print_action)
        menu.addAction(add_money_action)
        menu.addAction(pay_percent_action)
        menu.addAction(closing_contract_action)
        menu.addAction(add_to_the_black_list_action)
        # menu.addAction(transfer_to_selling_items_action)
        # menu.addAction(got_back_from_police_action)

        menu.exec_(self.table.viewport().mapToGlobal(position))


    # Right click menu functions
    # Print function menu
    def print_selected_row_page(self):
        selected = self.table.selectionModel().selectedRows()
        if not selected:

            return

        row_index = selected[0].row()
        contract_id = self.model.data(self.model.index(row_index, self.model.fieldIndex("id")))
        name = self.model.data(self.model.index(row_index, self.model.fieldIndex("name_surname")))
        given_money = self.model.data(self.model.index(row_index, self.model.fieldIndex("given_money")))
        date_raw = self.model.data(self.model.index(row_index, self.model.fieldIndex("date")))
        id_number = self.model.data(self.model.index(row_index, self.model.fieldIndex("id_number")))
        item_name = self.model.data(self.model.index(row_index, self.model.fieldIndex("item_name")))
        model = self.model.data(self.model.index(row_index, self.model.fieldIndex("model")))
        imei = self.model.data(self.model.index(row_index, self.model.fieldIndex("imei")))
        comment = self.model.data(self.model.index(row_index, self.model.fieldIndex("comment")))
        trusted_person = self.model.data(self.model.index(row_index, self.model.fieldIndex("trusted_person")))
        tel_number = self.model.data(self.model.index(row_index, self.model.fieldIndex("tel_number")))
        dt = datetime.strptime(date_raw, "%Y-%m-%d %H:%M:%S")
        date = dt.strftime("%d-%m-%Y")

        replacements = {
            '{name_surname}': name or "",
            '{given_money}': str(given_money) if given_money is not None else "",
            '{date}': date or "",
            '{contract_id}': str(contract_id or ""),
            '{id_number}': id_number or "",
            '{IMEI}': imei or "",
            '{model}': model or "",
            '{item_name}': item_name or "",
            '{comment}': comment or "",
            '{trusted_person}': trusted_person or "",
            '{tel_number}': tel_number or "",
            '{organization_name}': getattr(self, "organisation", ""),
            '{operator_name}': getattr(self, "name_of_user", "")
        }

        def replace_in_paragraph(paragraph, replacements):
            full_text = ''.join(run.text for run in paragraph.runs)
            new_text = full_text
            for key, value in replacements.items():
                new_text = new_text.replace(key, str(value))

            if new_text != full_text:
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

        # Load the Word template
        doc = Document("Templates/contract_template.docx")

        # Replace in normal paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)

        # Replace in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)

        # 3. Save new doc
        # Ensure folder exists
        output_dir = "GeneratedContracts"
        os.makedirs(output_dir, exist_ok=True)

        # Construct file name
        output_filename = f"contract_{contract_id}_{name}.docx"
        output_path = os.path.join(output_dir, output_filename)

        # Save document
        doc.save(output_path)

        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word_doc = word.Documents.Open(os.path.abspath(output_path))
            word_doc.PrintOut()  # Try to print
            word_doc.Close(False)  # Close document without saving
            word.Quit()
        except Exception as e:
            print("Printing failed:", e)
            print("Opening document instead...")
            try:
                os.startfile(output_path)  # Open in Word as fallback
            except Exception as open_error:
                print("Could not open Word document:", open_error)




    # Add money function menu
    def add_money_selected_row(self):
        row = self.table.currentIndex().row()
        if row < 0:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ ხელშეკრულება")
            return

        name_surname = self.model.data(self.model.index(row, self.model.fieldIndex("name_surname")))
        contract_id = self.model.data(self.model.index(row, self.model.fieldIndex("id")))
        self.open_add_money_window(contract_id, name_surname)


    def open_add_money_window(self, contract_id, name_surname):
        self.add_money_window = AddMoney(contract_id, name_surname, self.organisation)
        self.add_money_window.show()



    # Pay percent functions menu
    def pay_percent_selected_row(self):
        row = self.table.currentIndex().row()
        if row < 0:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ ხელშეკრულება")
            return

        contract_id = self.model.data(self.model.index(row, self.model.fieldIndex("id")))
        self.open_payment_window(contract_id)

    def open_payment_window(self, contract_id):
        self.payment_window = PaymentWindow(contract_id, self.organisation)
        self.payment_window.show()




    # Closing contract menu function
    def closing_contract_selected_row(self):
        row = self.table.currentIndex().row()
        if row < 0:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ ხელშეკრულება")
            return


        contract_id = self.model.data(self.model.index(row, self.model.fieldIndex("id")))
        name_surname = self.model.data(self.model.index(row, self.model.fieldIndex("name_surname")))
        principal_paid = self.model.data(self.model.index(row, self.model.fieldIndex("principal_paid")))
        percent_paid = self.model.data(self.model.index(row, self.model.fieldIndex("paid_percents")))
        given_money = self.model.data(self.model.index(row, self.model.fieldIndex("given_money")))
        principal_should_be_paid = self.model.data(self.model.index(row, self.model.fieldIndex("principal_should_be_paid")))

        self.confirm_window = PaymentConfirmWindow(contract_id, name_surname, principal_paid, percent_paid,
                                                   given_money, principal_should_be_paid)
        self.confirm_window.show()







    # Add to the black list function menu
    def add_to_the_black_list_selected_row(self):
        row = self.table.currentIndex().row()
        if row < 0:
            QMessageBox.warning(self, "შეცდომა", "გთხოვთ აირჩიოთ ხელშეკრულება")
            return

        name_surname_to_blk_list = self.model.data(self.model.index(row, self.model.fieldIndex("name_surname")))
        id_number_blk_list = self.model.data(self.model.index(row, self.model.fieldIndex("id_number")))
        tel_number_blk_list = self.model.data(self.model.index(row, self.model.fieldIndex("tel_number")))
        imei_blk_list = self.model.data(self.model.index(row, self.model.fieldIndex("imei")))

        try:
            conn = sqlite3.connect("Databases/black_list.db")  # Make sure this matches your DB
            cursor = conn.cursor()

            cursor.execute("""
                    INSERT INTO black_list (
                        name_surname, id_number, tel_number, imei
                    ) VALUES (?, ?, ?, ?)
                """, (
                name_surname_to_blk_list,
                id_number_blk_list,
                tel_number_blk_list,
                imei_blk_list
                ))

            conn.commit()
            QMessageBox.information(self, "წარმატება", "მონაცემები შენახულია")
        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", f"ვერ შევინახე მონაცემები:\n{e}")
        finally:
            conn.close()





    # # Transfer to selling items selected row function menu
    # def transfer_to_selling_items_selected_row(self):
    #     row = self.table.currentIndex().row()
    #     record_id = self.model.data(self.model.index(row, self.model.fieldIndex("id")))
    #     pass

    # # Items got back from the police function menu
    # def got_back_from_police_action_selected_row(self):
    #     row = self.table.currentIndex().row()
    #     record_id = self.model.data(self.model.index(row, self.model.fieldIndex("id")))
    #     pass




    # Initializing databases for money control window tables
    def initialize_contracts_database(self):
        conn = sqlite3.connect("Databases/contracts.db")
        cursor = conn.cursor()

        cursor.execute("""
                      CREATE TABLE IF NOT EXISTS contracts (
                          unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                          contract_id INTEGER,
                          contract_open_date TEXT,
                          first_percent_payment_date TEXT,
                          name_surname TEXT,
                          id_number TEXT,
                          tel_number TEXT,
                          item_name TEXT,
                          model TEXT,
                          IMEI TEXT,
                          trusted_person TEXT,
                          comment TEXT,
                          given_money INTEGER,
                          percent_day_quantity INTEGER,
                          first_added_percent REAL,
                          sum_of_principle_and_percent REAL GENERATED ALWAYS AS (given_money + first_added_percent) STORED,
                          office_mob_number TEXT
                      )
                  """)

        conn.commit()

        cursor.execute("DROP VIEW IF EXISTS contracts_view")

        cursor.execute("""
            CREATE VIEW contracts_view AS
            SELECT
                unique_id,
                contract_id,
                contract_open_date,
                first_percent_payment_date,
                name_surname,
                id_number,
                tel_number,
                item_name,
                model,
                IMEI,
                trusted_person,
                comment,
                given_money,
                percent_day_quantity,
                first_added_percent,
                (given_money + first_added_percent) AS sum_of_principle_and_percent,
                office_mob_number
            FROM contracts
        """)

        conn.commit()
        conn.close()



    def initialize_closed_contracts_database(self):
        conn = sqlite3.connect("Databases/closed_contracts.db")
        cursor = conn.cursor()

        cursor.execute("""
                      CREATE TABLE IF NOT EXISTS closed_contracts (
                          id INTEGER PRIMARY KEY,
                          contract_open_date TEXT,
                          name_surname TEXT,
                          id_number TEXT,
                          tel_number TEXT,
                          item_name TEXT,
                          model TEXT,
                          IMEI TEXT,
                          trusted_person TEXT,
                          comment TEXT,
                          percent REAL,
                          percent_day_quantity INTEGER,
                          given_money INTEGER,
                          additional_money INTEGER,
                          paid_principle REAL,
                          added_percents REAL,
                          paid_percents REAL,
                          status TEXT,
                          date_of_closing TEXT
                      )
                  """)

        conn.commit()
        conn.close()


    def initialize_active_contracts_database(self):
        conn = sqlite3.connect("Databases/active_contracts.db")
        cursor = conn.cursor()

        cursor.execute("""
            CREATE TABLE IF NOT EXISTS active_contracts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                date TEXT,
                days_after_C_O INTEGER NOT NULL DEFAULT 0,
                name_surname TEXT,
                id_number TEXT,
                tel_number TEXT,
                item_name TEXT,
                model TEXT,
                imei TEXT,
                type TEXT,
                trusted_person TEXT,
                comment TEXT,
                given_money REAL NOT NULL DEFAULT 0,
                percent REAL NOT NULL DEFAULT 0,
                day_quantity INTEGER NOT NULL DEFAULT 0,
                additional_amounts REAL NOT NULL DEFAULT 0,
                principal_paid REAL NOT NULL DEFAULT 0,
                principal_should_be_paid REAL GENERATED ALWAYS AS (
                    given_money + additional_amounts - principal_paid
                ) STORED,
                added_percents REAL NOT NULL DEFAULT 0,
                paid_percents REAL NOT NULL DEFAULT 0,
                percent_should_be_paid REAL GENERATED ALWAYS AS (
                    added_percents - paid_percents
                ) STORED,
                is_visible TEXT DEFAULT 'აქტიური'
            )
        """)

        conn.commit()

        # Create or replace the view that exposes generated columns explicitly
        cursor.execute("DROP VIEW IF EXISTS active_contracts_view")
        cursor.execute("""
               CREATE VIEW active_contracts_view AS
               SELECT
                   id,
                   date,
                   days_after_C_O,
                   name_surname,
                   id_number,
                   tel_number,
                   item_name,
                   model,
                   imei,
                   type,
                   trusted_person,
                   comment,
                   given_money,
                   percent,
                   day_quantity,
                   additional_amounts,
                   principal_paid,
                   (given_money + additional_amounts - principal_paid) AS principal_should_be_paid,
                   added_percents,
                   paid_percents,
                   (added_percents - paid_percents) AS percent_should_be_paid,
                   is_visible
               FROM active_contracts
           """)

        conn.commit()
        conn.close()

    @staticmethod
    def get_already_added_times(contract_id):
        conn_1 = sqlite3.connect("Databases/adding_percent_amount.db")
        cursor_1 = conn_1.cursor()
        cursor_1.execute("SELECT COUNT(*) FROM adding_percent_amount WHERE contract_id = ?",
                         (contract_id,))
        result = cursor_1.fetchone()[0]
        conn_1.close()
        return result

    def load_data(self):

        today = QDate.currentDate()

        # Open DB connection (adjust DB path as needed)
        conn = sqlite3.connect("Databases/active_contracts.db")
        cursor = conn.cursor()

        cursor.execute("SELECT * FROM active_contracts_view")  # Adjust table name & columns accordingly
        rows = cursor.fetchall()

        for row_index, row in enumerate(rows):
            try:

                contract_id = row[0]
                full_date_str = row[1]

                contract_datetime = QDateTime.fromString(full_date_str, "yyyy-MM-dd HH:mm:ss")
                if not contract_datetime.isValid():
                    raise ValueError(f"Invalid datetime format: {full_date_str}")
                contract_date = contract_datetime.date()

                name_surname = str(row[3])
                id_number = str(row[4])
                tel_number = str(row[5])
                item_name = str(row[6])
                model = str(row[7])
                imei_sn = str(row[8])

                # These could crash if values are None or strings
                day_quantity = int(row[14])
                percent = int(row[13])
                principal_should_be_paid = float(row[17])
                added_percents = float(row[18])
                paid_percents = float(row[19])
                percent_should_be_paid = float(row[20])

                days_after = contract_date.daysTo(today)
                new_added_percents = added_percents
                first_due_day = contract_date.addDays(day_quantity - 1)
                start_date = min(contract_date, first_due_day)

                if days_after >= day_quantity > 0 and principal_should_be_paid > 0 and percent > 0:
                    days_diff = start_date.daysTo(today)

                    periods_passed = days_after // day_quantity
                    total_expected_adds = periods_passed + 1


                    already_added_times = self.get_already_added_times(contract_id)
                    additions_needed = total_expected_adds - already_added_times

                    if additions_needed > 0:
                        one_period_amount = (principal_should_be_paid * percent) / 100
                        new_added_percents = added_percents
                        status_for_added_percent = "დარიცხული პროცენტი"

                        conn2 = sqlite3.connect("Databases/adding_percent_amount.db")
                        cursor2 = conn2.cursor()

                        for i in range(additions_needed):
                            new_added_percents += one_period_amount
                            cursor2.execute("""
                                INSERT INTO adding_percent_amount (
                                    contract_id, date_of_C_O, name_surname, id_number,
                                    tel_number, item_name, model, IMEI,
                                    date_of_percent_addition, percent_amount, status
                                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """, (
                                contract_id, full_date_str, name_surname, id_number,
                                tel_number, item_name, model, imei_sn,
                                QDateTime.currentDateTime().toString("yyyy-MM-dd HH:mm:ss"), one_period_amount,
                                status_for_added_percent
                            ))

                        conn2.commit()
                        conn2.close()

                # Update DB
                self.update_days_and_percents(contract_id, days_after, new_added_percents)


            except Exception as e:
                QMessageBox.critical(self, "Error Loading Data", f"Row {row_index + 1} caused an error:\n{e}")


    def update_days_and_percents(self, contract_id, days_after, new_added_percents):
        conn = sqlite3.connect("Databases/active_contracts.db")
        cursor = conn.cursor()
        cursor.execute("""
            UPDATE active_contracts
            SET days_after_C_O = ?, added_percents = ?
            WHERE id = ?
        """, (days_after, new_added_percents, contract_id))
        conn.commit()
        conn.close()


    def initialize_given_and_additional_database(self):
        conn = sqlite3.connect("Databases/given_and_additional_database.db")
        cursor = conn.cursor()

        cursor.execute("""
                      CREATE TABLE IF NOT EXISTS given_and_additional_database (
                          unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                          contract_id INTEGER,
                          date_of_outflow TEXT,
                          name_surname TEXT,
                          amount REAL,
                          status TEXT
                      )
                  """)

        conn.commit()
        conn.close()


    def initialize_paid_principle_and_paid_percentage_database(self):
        conn = sqlite3.connect("Databases/paid_principle_and_paid_percentage_database.db")
        cursor = conn.cursor()

        cursor.execute("""
                      CREATE TABLE IF NOT EXISTS paid_principle_and_paid_percentage_database (
                          unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                          contract_id INTEGER,
                          date_of_inflow TEXT,
                          name_surname TEXT,
                          amount REAL,
                          status TEXT
                      )
                  """)

        conn.commit()
        conn.close()

    def initialize_paid_principle_registry_database(self):
        conn = sqlite3.connect("Databases/paid_principle_registry.db")
        cursor = conn.cursor()

        cursor.execute("""
                      CREATE TABLE IF NOT EXISTS paid_principle_registry (
                          unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                          contract_id INTEGER,
                          date_of_C_O TEXT,
                          name_surname TEXT,
                          tel_number TEXT,
                          id_number TEXT,
                          item_name TEXT,
                          model TEXT,
                          IMEI TEXT,
                          given_money INTEGER,
                          date_of_payment TEXT,
                          payment_amount REAL,
                          status TEXT
                      )
                  """)

        conn.commit()
        conn.close()


    def initialize_outflow_order_database(self):
        conn = sqlite3.connect("Databases/outflow_order.db")
        cursor = conn.cursor()

        cursor.execute("""
                      CREATE TABLE IF NOT EXISTS outflow_order (
                          unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                          contract_id INTEGER,
                          name_surname TEXT,
                          tel_number TEXT,
                          amount REAL,
                          date TEXT,
                          status TEXT
                      )
                  """)

        conn.commit()
        conn.close()


    def initialize_outflow_in_registry_database(self):
        conn = sqlite3.connect("Databases/outflow_in_registry.db")
        cursor = conn.cursor()

        cursor.execute("""
                      CREATE TABLE IF NOT EXISTS outflow_in_registry (
                          unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                          contract_id INTEGER,
                          date_of_C_O TEXT,
                          name_surname TEXT,
                          tel_number TEXT,
                          id_number TEXT,
                          item_name TEXT,
                          model TEXT,
                          IMEI TEXT,
                          given_money INTEGER,
                          date_of_addition TEXT,
                          additional_amount INTEGER,
                          status TEXT
                      )
                  """)

        conn.commit()
        conn.close()

    def initialize_adding_percent_amount_database(self):
        conn = sqlite3.connect("Databases/adding_percent_amount.db")
        cursor = conn.cursor()

        cursor.execute("""
                              CREATE TABLE IF NOT EXISTS adding_percent_amount (
                                  unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                                  contract_id INTEGER,
                                  date_of_C_O TEXT,
                                  name_surname TEXT,
                                  tel_number TEXT,
                                  id_number TEXT,
                                  item_name TEXT,
                                  model TEXT,
                                  IMEI TEXT,
                                  date_of_percent_addition TEXT,
                                  percent_amount INTEGER,
                                  status TEXT
                              )
                          """)


        conn.commit()
        conn.close()

    def initialize_paid_percent_amount_database(self):
        conn = sqlite3.connect("Databases/paid_percent_amount.db")
        cursor = conn.cursor()

        cursor.execute("""
                              CREATE TABLE IF NOT EXISTS paid_percent_amount (
                                  unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                                  contract_id INTEGER,
                                  date_of_C_O TEXT,
                                  name_surname TEXT,
                                  tel_number TEXT,
                                  id_number TEXT,
                                  item_name TEXT,
                                  model TEXT,
                                  IMEI TEXT,
                                  set_date TEXT,
                                  date_of_percent_addition TEXT,
                                  paid_amount INTEGER,
                                  status TEXT
                              )
                          """)

        conn.commit()
        conn.close()


    def initialize_inflow_order_only_principal_amount_database(self):
        conn = sqlite3.connect("Databases/inflow_order_only_principal_amount.db")
        cursor = conn.cursor()

        cursor.execute("""
                              CREATE TABLE IF NOT EXISTS inflow_order_only_principal_amount (
                                  unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                                  contract_id INTEGER,
                                  name_surname TEXT,
                                  principle_paid_amount REAL,
                                  payment_date TEXT,
                                  sum_of_money_paid REAL
                              )
                          """)

        conn.commit()
        conn.close()


    def initialize_inflow_order_only_percent_amount_database(self):
        conn = sqlite3.connect("Databases/inflow_order_only_percent_amount.db")
        cursor = conn.cursor()

        cursor.execute("""
                              CREATE TABLE IF NOT EXISTS inflow_order_only_percent_amount (
                                  unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                                  contract_id INTEGER,
                                  name_surname TEXT,
                                  payment_date TEXT,
                                  set_date TEXT,
                                  percent_paid_amount INTEGER,
                                  sum_of_money_paid REAL
                              )
                          """)

        conn.commit()
        conn.close()


    def initialize_blk_list_database(self):
        conn = sqlite3.connect("Databases/black_list.db")
        cursor = conn.cursor()

        cursor.execute("""
                   CREATE TABLE IF NOT EXISTS black_list (
                       id INTEGER PRIMARY KEY AUTOINCREMENT,
                       name_surname TEXT,
                       id_number TEXT,
                       tel_number TEXT,
                       imei TEXT
                   )
               """)

        conn.commit()
        conn.close()

    def initialize_inflow_order_both_database(self):
        conn = sqlite3.connect("Databases/inflow_order_both.db")
        cursor = conn.cursor()

        cursor.execute("""
                CREATE TABLE IF NOT EXISTS inflow_order_both (
                    unique_id INTEGER PRIMARY KEY AUTOINCREMENT,
                    contract_id INTEGER,
                    name_surname TEXT,
                    payment_date TEXT,
                    principle_paid_amount REAL NOT NULL DEFAULT 0,
                    percent_paid_amount REAL NOT NULL DEFAULT 0,
                    sum_of_money_paid REAL GENERATED ALWAYS AS (
                        principle_paid_amount + percent_paid_amount
                    ) STORED
                )
            """)

        conn.commit()

        # Create a view for the table
        cursor.execute("DROP VIEW IF EXISTS inflow_order_both_view")

        cursor.execute("""
                CREATE VIEW inflow_order_both_view AS
                SELECT
                    unique_id,
                    contract_id,
                    name_surname,
                    payment_date,
                    principle_paid_amount,
                    percent_paid_amount,
                    sum_of_money_paid
                FROM inflow_order_both
            """)

        conn.commit()
        conn.close()

    def open_detail_window(self, index):
        row = index.row()
        contract_id = self.model.data(self.model.index(row, self.model.fieldIndex("id")))
        name_surname = self.model.data(self.model.index(row, self.model.fieldIndex("name_surname")))
        item_name = self.model.data(self.model.index(row, self.model.fieldIndex("item_name")))

        # Example: open a custom window and pass the contract ID
        self.detail_window = DetailWindow(contract_id, name_surname, item_name)
        self.detail_window.show()

    def update_summary_footer(self):
        sum_principal = 0.0
        sum_added_percents = 0.0

        for row in range(self.model.rowCount()):
            try:
                principal_index = self.model.index(row, self.model.fieldIndex("principal_should_be_paid"))
                percent_index = self.model.index(row, self.model.fieldIndex("added_percents"))

                principal_value = float(self.model.data(principal_index) or 0)
                percent_value = float(self.model.data(percent_index) or 0)

                sum_principal += principal_value
                sum_added_percents += percent_value
            except:
                continue

        self.total_label.setText(
            f"დაბანდებული თანხა (ძირი): {sum_principal:.2f} ₾     |     სულ დარიცხული პროცენტების ჯამი: {sum_added_percents:.2f} ₾"
        )