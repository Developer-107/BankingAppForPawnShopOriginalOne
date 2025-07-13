import os
import sqlite3
from datetime import datetime

import win32com.client
from docx import Document

from PyQt5.QtCore import QDate, QDateTime
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import QWidget, QGridLayout, QLabel, QLineEdit, QDateEdit, QComboBox, QPushButton, QMessageBox, \
    QDateTimeEdit


class PaymentWindow(QWidget):
    def __init__(self, contract_id, organisation):
        super().__init__()
        self.contract_id = contract_id
        self.setWindowTitle("პროცენტის ან ძირი თანხის გადახდა")
        self.setWindowIcon(QIcon("Icons/percent_payment_icon.png"))
        self.setFixedSize(400, 250)
        self.organisation = organisation

        layout = QGridLayout()

        # Contract ID (readonly)
        layout.addWidget(QLabel("ხელშეკრულების №:"), 0, 0)
        self.contract_id_box = QLineEdit(str(contract_id))
        self.contract_id_box.setReadOnly(True)
        layout.addWidget(self.contract_id_box, 0, 1)

        # Payment date
        layout.addWidget(QLabel("პროცენტის თარიღი:"), 1, 0)
        self.set_date = QDateTimeEdit(QDateTime.currentDateTime())
        self.set_date.setDisplayFormat("yyyy-MM-dd HH:mm:ss")
        self.set_date.setCalendarPopup(True)
        layout.addWidget(self.set_date, 1, 1)

        # Payment amount inputs
        layout.addWidget(QLabel("პროცენტის გადახდა:"), 2, 0)
        self.payed_percent_amount = QLineEdit()
        layout.addWidget(self.payed_percent_amount, 2, 1)

        layout.addWidget(QLabel("ძირის გადახდა:"), 3, 0)
        self.amount_input = QLineEdit()
        layout.addWidget(self.amount_input, 3, 1)

        # Buttons
        save_button = QPushButton("შენახვა")
        save_button.clicked.connect(self.save_the_payment)
        cancel_button = QPushButton("დახურვა")
        cancel_button.clicked.connect(self.close)
        layout.addWidget(save_button, 4, 0)
        layout.addWidget(cancel_button, 4, 1)

        self.setLayout(layout)

    def save_the_payment(self):

        status_for_payed_percent_money = "გადახდილი პროცენტი"
        status_for_payed_ground_money = "გადახდილი ძირი თანხა"
        payment_date = QDateTime.currentDateTime().toString("yyyy-MM-dd HH:mm:ss")
        set_date = self.set_date.dateTime().toString("yyyy-MM-dd HH:mm:ss")
        contract_id = self.contract_id_box.text()




        try:
            # Step 1: Get the id_number from the original active_contracts table
            source_conn = sqlite3.connect("Databases/active_contracts.db")
            source_cursor = source_conn.cursor()
            source_cursor.execute("""SELECT id_number, name_surname, principal_paid, date, item_name, imei, 
                                        tel_number, model, given_money, paid_percents 
                                        FROM active_contracts WHERE id = ?""", (contract_id,))
            result = source_cursor.fetchone()
            source_conn.close()

            if not result:
                QMessageBox.warning(self, "შეცდომა", "მითითებული ID-ით ჩანაწერი ვერ მოიძებნა active_contracts ბაზაში.")
                return

            id_number_from_contracts = result[0]
            name_from_contracts = result[1]
            old_principal_paid = result[2]
            date_of_c_o = result[3]
            item_name = result[4]
            imei = result[5]
            tel_number = result[6]
            model = result[7]
            given_money = result[8]
            old_percent_paid = result[9]


            if self.payed_percent_amount.text().strip():
                conn = sqlite3.connect("Databases/active_contracts.db")
                cursor = conn.cursor()

                new_percent_amount = old_percent_paid + int(self.payed_percent_amount.text())

                cursor.execute("""
                                    UPDATE active_contracts
                                    SET paid_percents = ?
                                    WHERE id = ?
                                """, (new_percent_amount, contract_id))

                conn.commit()
                conn.close()

                conn = sqlite3.connect("Databases/paid_percent_amount.db")
                cursor = conn.cursor()

                cursor.execute("""
                                               INSERT INTO paid_percent_amount (
                                                   contract_id, date_of_C_O, name_surname, tel_number, id_number, item_name,
                                                   model, IMEI, date_of_percent_addition, paid_amount, status, set_date
                                                   ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                           """, (
                    contract_id,
                    date_of_c_o,
                    name_from_contracts,
                    tel_number,
                    id_number_from_contracts,
                    item_name,
                    model,
                    imei,
                    payment_date,
                    int(self.payed_percent_amount.text()),
                    status_for_payed_percent_money,
                    set_date
                ))
                conn.commit()
                conn.close()

                conn = sqlite3.connect("Databases/paid_principle_and_paid_percentage_database.db")
                cursor = conn.cursor()

                cursor.execute("""
                                    INSERT INTO paid_principle_and_paid_percentage_database (
                                    contract_id, date_of_inflow, name_surname, amount, status
                                    ) VALUES (?, ?, ?, ?, ?)
                               """, (
                    contract_id,
                    payment_date,
                    name_from_contracts,
                    int(self.payed_percent_amount.text()),
                    status_for_payed_percent_money
                ))
                conn.commit()
                conn.close()

                conn = sqlite3.connect("Databases/inflow_order_only_percent_amount.db")
                cursor = conn.cursor()

                cursor.execute("""
                                      INSERT INTO inflow_order_only_percent_amount (
                                      contract_id, payment_date, name_surname, percent_paid_amount, sum_of_money_paid, 
                                      set_date
                                      ) VALUES (?, ?, ?, ?, ?, ?)
                               """, (
                    contract_id,
                    payment_date,
                    name_from_contracts,
                    int(self.payed_percent_amount.text()),
                    int(self.payed_percent_amount.text()),
                    set_date
                ))

                self.unique_id = cursor.lastrowid

                conn.commit()
                conn.close()


            if self.amount_input.text().strip():
                conn = sqlite3.connect("Databases/active_contracts.db")
                cursor = conn.cursor()

                new_principal_amount = old_principal_paid + int(self.amount_input.text())

                cursor.execute("""
                    UPDATE active_contracts
                    SET principal_paid = ?
                    WHERE id = ?
                """, (new_principal_amount, contract_id))

                conn.commit()
                conn.close()

                conn = sqlite3.connect("Databases/paid_principle_registry.db")
                cursor = conn.cursor()

                cursor.execute("""
                                INSERT INTO paid_principle_registry (
                                    contract_id, date_of_C_O, name_surname, tel_number, id_number, item_name,
                                    model, IMEI, given_money, date_of_payment, payment_amount, status
                                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                            """, (
                    contract_id,
                    date_of_c_o,
                    name_from_contracts,
                    tel_number,
                    id_number_from_contracts,
                    item_name,
                    model,
                    imei,
                    given_money,
                    payment_date,
                    int(self.amount_input.text()),
                    status_for_payed_ground_money
                ))
                conn.commit()
                conn.close()

                conn = sqlite3.connect("Databases/paid_principle_and_paid_percentage_database.db")
                cursor = conn.cursor()

                cursor.execute("""
                                               INSERT INTO paid_principle_and_paid_percentage_database (
                                                   contract_id, date_of_inflow, name_surname, amount, status
                                                   ) VALUES (?, ?, ?, ?, ?)
                                           """, (
                    contract_id,
                    payment_date,
                    name_from_contracts,
                    int(self.amount_input.text()),
                    status_for_payed_ground_money
                ))
                conn.commit()
                conn.close()

                conn = sqlite3.connect("Databases/inflow_order_only_principal_amount.db")
                cursor = conn.cursor()

                cursor.execute("""
                                INSERT INTO inflow_order_only_principal_amount (
                                contract_id, payment_date, name_surname, principle_paid_amount, sum_of_money_paid
                                ) VALUES (?, ?, ?, ?, ?)
                    """, (
                    contract_id,
                    payment_date,
                    name_from_contracts,
                    int(self.amount_input.text()),
                    int(self.amount_input.text())
                ))

                self.unique_id = cursor.lastrowid

                conn.commit()
                conn.close()

            if self.payed_percent_amount.text().strip() and self.amount_input.text().strip():

                conn = sqlite3.connect("Databases/inflow_order_both.db")
                cursor = conn.cursor()

                cursor.execute("""
                                                       INSERT INTO inflow_order_both (
                                                       contract_id, payment_date, name_surname, principle_paid_amount,
                                                       percent_paid_amount
                                                       ) VALUES (?, ?, ?, ?, ?)
                                               """, (
                    contract_id,
                    payment_date,
                    name_from_contracts,
                    int(self.amount_input.text()),
                    int(self.payed_percent_amount.text())
                ))

                self.unique_id = cursor.lastrowid

                conn.commit()
                conn.close()

            if not self.payed_percent_amount.text().strip() and not self.amount_input.text().strip():
                QMessageBox.warning(self, "შეცდომა", "გთხოვთ მიუთითოთ მაინც ერთი გადახდის თანხა.")
                return


            QMessageBox.information(self, "წარმატება", "მონაცემები შენახულია")
            self.close()

            reply = QMessageBox.question(
                self,
                "ბეჭდვა",
                "გსურთ დოკუმენტის ამობეჭდვა?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )

            if reply == QMessageBox.Yes:
                try:
                    dt = datetime.strptime(payment_date, "%Y-%m-%d %H:%M:%S")
                    date = dt.strftime("%d-%m-%Y")
                    amount_input_value = int(self.amount_input.text()) if self.amount_input.text().strip().isdigit() else 0
                    percent_input_value = int(
                        self.payed_percent_amount.text()) if self.payed_percent_amount.text().strip().isdigit() else 0
                    sum_of_paid_amount = amount_input_value + percent_input_value


                    replacements = {
                        '{name_surname}': name_from_contracts or "",
                        '{paid_principle}': str(amount_input_value),
                        '{paid_percent}': str(percent_input_value),
                        '{sum}': str(sum_of_paid_amount),
                        '{date_of_payment}': date or "",
                        '{contract_id}': str(contract_id or ""),
                        '{unique_id}': self.unique_id or "",
                        '{organization_name}': getattr(self, "organisation", ""),
                    }

                    def replace_in_paragraph(paragraph, replacements):
                        full_text = ''.join(run.text for run in paragraph.runs)
                        new_text = full_text
                        for key, value in replacements.items():
                            new_text = new_text.replace(key, str(value))

                        if new_text != full_text:
                            for run in paragraph.runs:
                                run.text = ''
                            if paragraph.runs:
                                paragraph.runs[0].text = new_text
                            else:
                                paragraph.add_run(new_text)

                    # Load the Word template
                    doc = Document("Templates/inflow_template.docx")

                    # Replace in normal paragraphs
                    for paragraph in doc.paragraphs:
                        replace_in_paragraph(paragraph, replacements)

                    # Replace in table cells
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    replace_in_paragraph(paragraph, replacements)

                    # 3. Save new doc
                    # Ensure folder exists
                    output_dir = "GeneratedContracts"
                    os.makedirs(output_dir, exist_ok=True)

                    # Construct file name
                    output_filename = f"inflow_order_both_{self.unique_id}_{contract_id}_{name_from_contracts}.docx"
                    output_path = os.path.join(output_dir, output_filename)

                    # Save document
                    doc.save(output_path)

                    try:
                        word = win32com.client.Dispatch("Word.Application")
                        word.Visible = False
                        word_doc = word.Documents.Open(os.path.abspath(output_path))
                        word_doc.PrintOut()  # Try to print
                        word_doc.Close(False)  # Close document without saving
                        word.Quit()
                    except Exception as e:
                        print("Printing failed:", e)
                        print("Opening document instead...")
                        try:
                            os.startfile(output_path)  # Open in Word as fallback
                        except Exception as open_error:
                            print("Could not open Word document:", open_error)
                except Exception as e:
                    QMessageBox.critical(self, "შეცდომა", f"ამობეჭდვა არ მოხერხდა:\n{e}")


        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", f"ვერ შევინახე მონაცემები:\n{e}")
