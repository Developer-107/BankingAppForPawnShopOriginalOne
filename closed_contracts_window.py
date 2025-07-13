import sqlite3
import sys
from datetime import datetime, timedelta
import win32com.client
from PyQt5.QtCore import Qt, QSize, QDate
from PyQt5.QtGui import QIcon, QTextDocument
from PyQt5.QtPrintSupport import QPrinter, QPrintDialog
from PyQt5.QtSql import QSqlDatabase, QSqlTableModel
from PyQt5.QtWidgets import QWidget, QGridLayout, QGroupBox, QToolButton, QCheckBox, QLabel, QRadioButton, QButtonGroup, \
    QLineEdit, QDateEdit, QPushButton, QTableView, QAbstractItemView, QMessageBox, QMenu, QAction
from docx import Document
import tempfile
import pandas as pd
import os
import subprocess
from PyQt5.QtWidgets import QMessageBox

from add_money_window import AddMoney
from open_edit_window import EditWindow
from open_add_window import AddWindow
from payment_confirm_window import PaymentConfirmWindow
from payment_window import PaymentWindow
from utils import resource_path


class ClosedContracts(QWidget):
    def __init__(self, role, name_of_user, organisation):
        super().__init__()
        self.setWindowTitle("დახურული ხელშეკრულებები")
        self.setWindowIcon(QIcon(resource_path("Icons/closed_contracts.png")))
        self.resize(1400, 800)
        self.role = role
        self.organisation = organisation
        self.name_of_user = name_of_user

        layout = QGridLayout()

        # --------------------------------------------Box1-----------------------------------------------------
        box1 = QGroupBox("ნავიგაცია")
        box1.setFixedWidth(275)
        box1.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout = QGridLayout()


        # Export
        export = QToolButton()
        export.setText(" ექსპორტი ")
        export.setIcon(QIcon(resource_path("Icons/excel_icon.png")))
        export.setIconSize(QSize(37, 40))
        export.setToolButtonStyle(Qt.ToolButtonTextBesideIcon)
        export.setStyleSheet("font-size: 16px;")
        # export.clicked.connect(self.export_to_excel)

        # Checkbox
        # all_together = QCheckBox("ყველა ერთად")

        # BlankUnderTheCheckBox
        # blank_under_the_check_box = QLabel("")
        # blank_under_the_check_box.setStyleSheet("border: 1px solid gray; padding: 5px;")
        # blank_under_the_check_box.setFixedHeight(30)


        # # Box1 widgets
        # box_layout.addWidget(add_item, 0, 0)
        # box_layout.addWidget(edit_item, 0, 2)
        # box_layout.addWidget(choosing_the_field, 0, 2)
        box_layout.addWidget(export, 0, 3)
        # box_layout.addWidget(all_together, 1, 1)
        # box_layout.addWidget(blank_under_the_check_box, 1, 2)

        # Placing box1
        layout.addWidget(box1, 0, 0, 1, 1)
        box1.setLayout(box_layout)

        # --------------------------------------------Box2-----------------------------------------------------
        box2 = QGroupBox("ძებნა")
        box2.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout2 = QGridLayout()

        self.search_input = QLineEdit()
        self.search_input.setPlaceholderText("მოძებნე აქ")
        self.search_input.textChanged.connect(self.apply_text_filter)

        # Radio Button
        self.contract_radio = QRadioButton("ხელშეკრულების N")
        self.name_radio = QRadioButton("სახელი და გვარი")
        self.id_radio = QRadioButton("პირადი N")
        self.model_radio = QRadioButton("მოდელი")
        self.tel_radio = QRadioButton("ტელეფონის ნომერი")
        self.imei_radio = QRadioButton("IMEI")

        # Grouping buttons
        button_group = QButtonGroup()
        button_group.addButton(self.contract_radio)
        button_group.addButton(self.name_radio)
        button_group.addButton(self.id_radio)
        button_group.addButton(self.model_radio)
        button_group.addButton(self.tel_radio)
        button_group.addButton(self.imei_radio)

        # Making Widgets
        box_layout2.addWidget(self.contract_radio)
        box_layout2.addWidget(self.name_radio)
        box_layout2.addWidget(self.id_radio)
        box_layout2.addWidget(self.model_radio)
        box_layout2.addWidget(self.tel_radio)
        box_layout2.addWidget(self.imei_radio)
        box_layout2.addWidget(self.search_input, 2, 1)



        layout.addWidget(box2, 0, 1, 1, 1)
        box2.setLayout(box_layout2)

        # --------------------------------------------Box3-----------------------------------------------------
        box3 = QGroupBox("თარიღის მიხედვით გაფილტვრა")
        box3.setStyleSheet("QGroupBox {font-style: italic; font-size: 10pt; }")
        box_layout3 = QGridLayout()

        self.from_date = QDateEdit()
        self.from_date.setCalendarPopup(True)
        self.from_date.setDate(QDate.currentDate())  # Today

        self.to_date = QDateEdit()
        self.to_date.setCalendarPopup(True)
        self.to_date.setDate(QDate.currentDate().addDays(1))  # Tomorrow

        self.contract_date_radio = QRadioButton("გაფორმების თარიღით")
        self.contract_date_radio.setChecked(True)  # Default
        self.closing_date_radio = QRadioButton("დახურვის თარიღით")

        box_layout3.addWidget(self.contract_date_radio, 0, 0, 1, 2)
        box_layout3.addWidget(self.closing_date_radio, 1, 0, 1, 2)

        box_layout3.addWidget(QLabel("დან თარიღი:"), 2, 0)
        box_layout3.addWidget(self.from_date, 2, 1)
        box_layout3.addWidget(QLabel("მდე თარიღი:"), 3, 0)
        box_layout3.addWidget(self.to_date, 3, 1)

        filter_button = QPushButton("ძებნა")
        filter_button.clicked.connect(self.search_by_date)
        refresh_button = QPushButton("განახლება")
        refresh_button.clicked.connect(self.refresh_table)

        box_layout3.addWidget(refresh_button, 4, 0)
        box_layout3.addWidget(filter_button, 4, 1)

        layout.addWidget(box3, 0, 2, 1, 1)
        box3.setLayout(box_layout3)


        # --------------------------------------------Table-----------------------------------------------------
        self.db = QSqlDatabase.addDatabase("QSQLITE")
        self.db.setDatabaseName(resource_path("Databases/closed_contracts.db"))
        if not self.db.open():
            raise Exception("ბაზასთან კავშირი ვერ მოხერხდა")

        self.table = QTableView()
        self.model = QSqlTableModel(self, self.db)
        self.model.setTable("closed_contracts")
        self.model.select()

        # Block to rename columns
        record = self.model.record()
        column_indices = {record.field(i).name(): i for i in range(record.count())}

        column_labels = {
            "id": "ხელშეკრულების N",
            "contract_open_date": "გაფორმების თარიღი",
            "name_surname": "სახელი და გვარი",
            "id_number": "პირადი ნომერი",
            "tel_number": "ტელეფონის ნომერი",
            "item_name": "ნივთის დასახელება",
            "model": "მოდელი",
            "trusted_person": "მინდობილი პირი",
            "comment": "კომენტარი",
            "given_money": "გაცემული ძირი თანხა",
            "percent": "პროცენტი",
            "percent_day_quantity": "დღეების რაოდენობა",
            "additional_money": "დამატებული თანხები",
            "paid_principle": "გადახდილი ძირი თანხა",
            "added_percents": "დარიცხული პროცენტები",
            "paid_percents": "გადახდილი პროცენტები",
            "status": "სტატუსი",
            "date_of_closing": "ხელშეკრულების დახურვის თარიღი"
        }

        for name, label in column_labels.items():
            if name in column_indices:
                self.model.setHeaderData(column_indices[name], Qt.Horizontal, label)


        self.table.setModel(self.model)
        # Make header text bold
        header = self.table.horizontalHeader()
        font = header.font()
        font.setBold(True)
        header.setFont(font)
        header.setStyleSheet("""
                    QHeaderView::section {
                        padding: 4px 8px;
                    }
                """)
        self.table.setEditTriggers(QAbstractItemView.NoEditTriggers)  # Read-only table
        self.table.setSelectionBehavior(QTableView.SelectRows)
        self.table.setSelectionMode(QTableView.SingleSelection)
        self.table.setContextMenuPolicy(Qt.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_closed_contracts_context_menu)

        self.table.resizeColumnsToContents()
        layout.addWidget(self.table, 1, 0, 4, 5)


        # --------------------------------------------Layout-----------------------------------------------------
        self.setLayout(layout)


        # --------------------------------------------Functions-----------------------------------------------------

    def show_closed_contracts_context_menu(self, position):
        index = self.table.indexAt(position)
        if not index.isValid():
            return

        menu = QMenu()

        print_action = QAction(" ამობეჭდვა ", self)
        print_action.setIcon(QIcon(resource_path("Icons/printer_icon.png")))
        print_action.triggered.connect(self.print_closed_contract_selected_row)
        menu.addAction(print_action)

        return_action = QAction(" დააბრუნე აქტიურში ", self)
        return_action.setIcon(QIcon(resource_path("Icons/return_icon.png")))
        return_action.triggered.connect(self.return_closed_contract_to_active)
        menu.addAction(return_action)

        menu.exec_(self.table.viewport().mapToGlobal(position))

    def print_closed_contract_selected_row(self):
        selected = self.table.selectionModel().selectedRows()
        if not selected:
            print("No row selected.")
            return

        row_index = selected[0].row()
        contract_id = self.model.data(self.model.index(row_index, self.model.fieldIndex("id")))
        name = self.model.data(self.model.index(row_index, self.model.fieldIndex("name_surname")))
        given_money = self.model.data(self.model.index(row_index, self.model.fieldIndex("given_money")))
        date_raw = self.model.data(self.model.index(row_index, self.model.fieldIndex("contract_open_date")))
        id_number = self.model.data(self.model.index(row_index, self.model.fieldIndex("id_number")))
        item_name = self.model.data(self.model.index(row_index, self.model.fieldIndex("item_name")))
        model = self.model.data(self.model.index(row_index, self.model.fieldIndex("model")))
        imei = self.model.data(self.model.index(row_index, self.model.fieldIndex("IMEI")))
        comment = self.model.data(self.model.index(row_index, self.model.fieldIndex("comment")))
        trusted_person = self.model.data(self.model.index(row_index, self.model.fieldIndex("trusted_person")))
        tel_number = self.model.data(self.model.index(row_index, self.model.fieldIndex("tel_number")))
        dt = datetime.strptime(date_raw, "%Y-%m-%d %H:%M:%S")
        date = dt.strftime("%d-%m-%Y")

        replacements = {
            '{name_surname}': name or "",
            '{given_money}': str(given_money) if given_money is not None else "",
            '{date}': date or "",
            '{contract_id}': str(contract_id or ""),
            '{id_number}': id_number or "",
            '{IMEI}': imei or "",
            '{model}': model or "",
            '{item_name}': item_name or "",
            '{comment}': comment or "",
            '{trusted_person}': trusted_person or "",
            '{tel_number}': tel_number or "",
            '{organization_name}': getattr(self, "organisation", ""),
            '{operator_name}': getattr(self, "name_of_user", "")
        }

        def replace_in_paragraph(paragraph, replacements):
            full_text = ''.join(run.text for run in paragraph.runs)
            new_text = full_text
            for key, value in replacements.items():
                new_text = new_text.replace(key, str(value))

            if new_text != full_text:
                for run in paragraph.runs:
                    run.text = ''
                if paragraph.runs:
                    paragraph.runs[0].text = new_text
                else:
                    paragraph.add_run(new_text)

        # Load the Word template
        doc = Document(resource_path("Templates/contract_template.docx"))

        # Replace in normal paragraphs
        for paragraph in doc.paragraphs:
            replace_in_paragraph(paragraph, replacements)

        # Replace in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        replace_in_paragraph(paragraph, replacements)

        # 3. Save new doc
        # Ensure folder exists
        output_dir = "GeneratedContracts"
        os.makedirs(output_dir, exist_ok=True)

        # Construct file name
        output_filename = f"contract_{contract_id}_{name}.docx"
        output_path = os.path.join(output_dir, output_filename)

        # Save document
        doc.save(output_path)

        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word_doc = word.Documents.Open(os.path.abspath(output_path))
            word_doc.PrintOut()  # Try to print
            word_doc.Close(False)  # Close document without saving
            word.Quit()
        except Exception as e:
            print("Printing failed:", e)
            print("Opening document instead...")
            try:
                os.startfile(output_path)  # Open in Word as fallback
            except Exception as open_error:
                print("Could not open Word document:", open_error)

    def return_closed_contract_to_active(self):
        selected = self.table.selectionModel().selectedRows()
        if not selected:
            QMessageBox.warning(self, "შეცდომა", "არ არის არჩეული რიგი")
            return

        row_index = selected[0].row()
        contract_id = self.model.data(self.model.index(row_index, self.model.fieldIndex("id")))

        try:
            # 1. Delete from closed_contracts DB
            conn_closed = sqlite3.connect(resource_path("Databases/closed_contracts.db"))
            cursor_closed = conn_closed.cursor()
            cursor_closed.execute("DELETE FROM closed_contracts WHERE id = ?", (contract_id,))
            conn_closed.commit()
            conn_closed.close()

            # 2. Update is_visible = 'აქტიური' in active_contracts DB
            conn_active = sqlite3.connect(resource_path("Databases/active_contracts.db"))
            cursor_active = conn_active.cursor()
            cursor_active.execute("""
                UPDATE active_contracts
                SET is_visible = 'აქტიური'
                WHERE id = ?
            """, (contract_id,))
            conn_active.commit()
            conn_active.close()

            QMessageBox.information(self, "წარმატება", "ხელშეკრულება დაბრუნდა აქტიურებში")

            # Optional: Refresh the model/view after changes
            self.model.select()

        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", f"მოცემულობის დამუშავებისას მოხდა შეცდომა:\n{e}")


    def refresh_table(self):
        self.model.setFilter("")  # Clears filter
        self.model.select()  # This reloads the data from DB

    def search_by_date(self):
        from_date_str = self.from_date.date().toString("dd.MM.yyyy")
        to_date_str = self.to_date.date().toString("dd.MM.yyyy")

        if self.contract_date_radio.isChecked():
            date_column = "date"
        elif self.closing_date_radio.isChecked():
            date_column = "closing_date"
        else:
            QMessageBox.warning(self, "შეცდომა", "აირჩიეთ თარიღის ტიპი.")
            return

        # Assuming your table has a date column named 'date'
        filter_str = f"{date_column} >= '{from_date_str}' AND date <= '{to_date_str}'"
        self.model.setFilter(filter_str)
        self.model.select()

    def apply_text_filter(self, text):
        column = ""

        if self.contract_radio.isChecked():
            column = "id"
        elif self.name_radio.isChecked():
            column = "name_surname"
        elif self.id_radio.isChecked():
            column = "id_number"
        elif self.model_radio.isChecked():
            column = "model"
        elif self.tel_radio.isChecked():
            column = "tel_number"
        elif self.imei_radio.isChecked():
            column = "imei"

        if column:
            filter_str = f"{column} LIKE '%{text}%'"
            self.model.setFilter(filter_str)
        else:
            self.model.setFilter("")  # No filter if nothing selected

        self.model.select()



    def export_to_excel(self):

        row_count = self.model.rowCount()
        col_count = self.model.columnCount()

        headers = [self.model.headerData(col, Qt.Horizontal) for col in range(col_count)]
        data = [
            [self.model.data(self.model.index(row, col)) for col in range(col_count)]
            for row in range(row_count)
        ]

        df = pd.DataFrame(data, columns=headers)

        try:
            temp_path = os.path.join(tempfile.gettempdir(), "temp_export.xlsx")
            df.to_excel(temp_path, index=False)

            # Open Excel file
            os.startfile(temp_path)  # Safer and native on Windows
        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", str(e))

