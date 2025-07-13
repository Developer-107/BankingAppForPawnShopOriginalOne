import os
import sqlite3
from datetime import datetime

import win32com.client
from docx import Document
from PyQt5.QtCore import QDate, QDateTime
from PyQt5.QtGui import QIcon
from PyQt5.QtWidgets import QWidget, QGridLayout, QLabel, QLineEdit, QDateEdit, QComboBox, QPushButton, QMessageBox


class AddMoney(QWidget):
    def __init__(self, contract_id, name_surname, organisation):
        super().__init__()
        self.contract_id = contract_id
        self.name_surname = name_surname
        self.setWindowTitle("თანხის დამატება")
        self.setWindowIcon(QIcon("Icons/add_money.png"))
        self.setFixedSize(400, 250)
        self.organisation = organisation

        layout = QGridLayout()

        # Contract ID (readonly)
        layout.addWidget(QLabel("ხელშეკრულების №:"), 0, 0)
        self.contract_id_box = QLineEdit(str(contract_id))
        self.contract_id_box.setReadOnly(True)
        layout.addWidget(self.contract_id_box, 0, 1)

        # Name (readonly)
        layout.addWidget(QLabel("სახელი და გვარი: "), 1, 0)
        self.name_box = QLineEdit(str(name_surname))
        self.name_box.setReadOnly(True)
        layout.addWidget(self.name_box, 1, 1)

        # Add money input
        layout.addWidget(QLabel("დამატებული თანხა: "), 2, 0)
        self.added_money_amount = QLineEdit()
        layout.addWidget(self.added_money_amount, 2, 1)

        # Buttons
        save_button = QPushButton("შენახვა")
        save_button.clicked.connect(self.save_payment)
        cancel_button = QPushButton("დახურვა")
        cancel_button.clicked.connect(self.close)
        layout.addWidget(save_button, 3, 0)
        layout.addWidget(cancel_button, 3, 1)

        self.setLayout(layout)


    def save_payment(self):

        status_for_added_money = "დამატებული"
        date_of_addition = QDateTime.currentDateTime().toString("yyyy-MM-dd HH:mm:ss")
        contract_id = self.contract_id_box.text()

        try:
            # Step 1: Get the id_number from the original active_contracts table
            source_conn = sqlite3.connect("Databases/active_contracts.db")
            source_cursor = source_conn.cursor()
            source_cursor.execute("""SELECT id_number, additional_amounts, date, item_name, 
                                                 model, imei, given_money, tel_number, percent
                                            FROM active_contracts WHERE id = ?""", (contract_id,))
            result = source_cursor.fetchone()

            if not result:
                QMessageBox.warning(self, "შეცდომა", "მითითებული ID-ით ჩანაწერი ვერ მოიძებნა contracts ბაზაში.")
                return

            id_number_from_contracts = result[0]
            additional_amounts = result[1]
            contract_open_date = result[2]
            item_name = result[3]
            model = result[4]
            imei = result[5]
            given_money = result[6]
            tel_number = result[7]
            percent = result[8]

            updated_additional_amount = float(additional_amounts) + int(self.added_money_amount.text())
            new_added_percents = (given_money + updated_additional_amount) * percent / 100
            source_cursor.execute("""
                    UPDATE active_contracts
                    SET additional_amounts = ?, added_percents = ?
                    WHERE id = ?
                    """, (updated_additional_amount, new_added_percents,contract_id,))

            source_conn.commit()
            source_conn.close()



            conn = sqlite3.connect("Databases/given_and_additional_database.db")  # Make sure this matches your DB
            cursor = conn.cursor()


            # Insert in given_and_additional_database database
            cursor.execute("""
                INSERT INTO given_and_additional_database (
                    contract_id, date_of_outflow, name_surname, amount, status
                ) VALUES (?, ?, ?, ?, ?)
            """, (
                self.contract_id_box.text(),
                date_of_addition,
                self.name_box.text(),
                int(self.added_money_amount.text()),
                status_for_added_money
                ))

            conn.commit()
            conn.close()

            conn = sqlite3.connect("Databases/outflow_order.db")  # Make sure this matches your DB
            cursor = conn.cursor()

            # Insert in outflow_order database
            cursor.execute("""
                            INSERT INTO outflow_order (
                                contract_id, date, name_surname, tel_number, amount, status
                            ) VALUES (?, ?, ?, ?, ?, ?)
                        """, (
                self.contract_id_box.text(),
                date_of_addition,
                self.name_box.text(),
                tel_number,
                int(self.added_money_amount.text()),
                status_for_added_money
            ))

            conn.commit()
            conn.close()

            conn = sqlite3.connect("Databases/outflow_in_registry.db")  # Make sure this matches your DB
            cursor = conn.cursor()

            # Insert in outflow_in_registry database
            cursor.execute("""
                            INSERT INTO outflow_in_registry (
                                contract_id, date_of_C_O, name_surname, tel_number, id_number, item_name, model, IMEI,
                                given_money, date_of_addition, additional_amount, status
                            ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                        """, (
                self.contract_id_box.text(),
                contract_open_date,
                self.name_box.text(),
                tel_number,
                id_number_from_contracts,
                item_name,
                model,
                imei,
                given_money,
                date_of_addition,
                int(self.added_money_amount.text()),
                status_for_added_money
            ))

            unique_id = cursor.lastrowid

            conn.commit()
            conn.close()


            QMessageBox.information(self, "წარმატება", "მონაცემები შენახულია")
            self.close()

            try:
                dt = datetime.strptime(date_of_addition, "%Y-%m-%d %H:%M:%S")
                date = dt.strftime("%d-%m-%Y")

                replacements = {
                    '{name_surname}': self.name_box.text() or "",
                    '{additional_amount}': str(int(self.added_money_amount.text())) if
                                    int(self.added_money_amount.text()) is not None else "",
                    '{date}': date or "",
                    '{contract_id}': str(contract_id or ""),
                    '{unique_id}': unique_id or "",
                    '{organization_name}': getattr(self, "organisation", ""),
                }

                def replace_in_paragraph(paragraph, replacements):
                    full_text = ''.join(run.text for run in paragraph.runs)
                    new_text = full_text
                    for key, value in replacements.items():
                        new_text = new_text.replace(key, str(value))

                    if new_text != full_text:
                        for run in paragraph.runs:
                            run.text = ''
                        if paragraph.runs:
                            paragraph.runs[0].text = new_text
                        else:
                            paragraph.add_run(new_text)

                # Load the Word template
                doc = Document("Templates/additional_money_template.docx")

                # Replace in normal paragraphs
                for paragraph in doc.paragraphs:
                    replace_in_paragraph(paragraph, replacements)

                # Replace in table cells
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                replace_in_paragraph(paragraph, replacements)

                # 3. Save new doc
                # Ensure folder exists
                output_dir = "GeneratedContracts"
                os.makedirs(output_dir, exist_ok=True)

                # Construct file name
                output_filename = f"outflow_order_{unique_id}_{contract_id}_{self.name_box.text()}.docx"
                output_path = os.path.join(output_dir, output_filename)

                # Save document
                doc.save(output_path)

                try:
                    word = win32com.client.Dispatch("Word.Application")
                    word.Visible = False
                    word_doc = word.Documents.Open(os.path.abspath(output_path))
                    word_doc.PrintOut()  # Try to print
                    word_doc.Close(False)  # Close document without saving
                    word.Quit()
                except Exception as e:
                    print("Printing failed:", e)
                    print("Opening document instead...")
                    try:
                        os.startfile(output_path)  # Open in Word as fallback
                    except Exception as open_error:
                        print("Could not open Word document:", open_error)
            except Exception as e:
                QMessageBox.critical(self, "შეცდომა", f"ამობეჭდვა ვერ მოოხერხდა:\n{e}")
        except Exception as e:
            QMessageBox.critical(self, "შეცდომა", f"ვერ შევინახე მონაცემები:\n{e}")
